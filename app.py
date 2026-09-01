import os
import json
import base64
import hmac
import math
import re
import textwrap
from urllib.parse import quote
from io import BytesIO
from datetime import date, timedelta

import requests

import streamlit as st
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    from PIL import Image, ImageDraw, ImageFont
except Exception:
    Image = ImageDraw = ImageFont = None

# =====================
# PAGE CONFIG
# =====================
st.set_page_config(
    page_title="EWMT AI Document System",
    page_icon="🏗️",
    layout="wide"
)

# =====================
# SESSION PAGE CONTROL
# =====================
PAGES = [
    "🏠 Dashboard",
    "📄 Method Statement",
    "🏗️ Lifting Plan",
    "⚠️ Risk Assessment Pro",
    "🧰 Lifting Gear Register",
    "👷 Worker Training Certificate",
    "⏰ Expiry Alerts",
    "⚙️ Settings"
]

if "page" not in st.session_state:
    st.session_state.page = "🏠 Dashboard"

# =====================
# OPENAI CLIENT
# =====================
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

# =====================
# VECTOR STORES
# =====================
MS_VECTOR_STORE_ID = "vs_69ecc533a1208191a8595c674753e99e"
RA_VECTOR_STORE_ID = "vs_69ecd191648481919d1d1d57f21264af"
LP_VECTOR_STORE_ID = "vs_69ecdc44d59081919fb10574510b7454"

# =====================
# PATHS + ASSETS
# =====================
BASE_DIR = os.path.dirname(__file__)
ASSET_DIR = os.path.join(BASE_DIR, "assets")

MS_TEMPLATE = os.path.join(BASE_DIR, "Templates", "Method of statement Template.docx")
RA_TEMPLATE = os.path.join(BASE_DIR, "Templates", "RA Template.docx")
LP_TEMPLATE = os.path.join(BASE_DIR, "Templates", "Lifting Plan Template.docx")


def image_to_base64(path):
    try:
        with open(path, "rb") as img_file:
            return base64.b64encode(img_file.read()).decode()
    except Exception:
        return ""


def asset_image(filename, fallback="banner.jpg"):
    path = os.path.join(ASSET_DIR, filename)

    if os.path.exists(path):
        return image_to_base64(path)

    fallback_path = os.path.join(ASSET_DIR, fallback)

    if os.path.exists(fallback_path):
        return image_to_base64(fallback_path)

    return ""


HEADER_IMAGE = asset_image("banner.jpg")
METHOD_IMAGE = asset_image("method_statement.jpg")
LIFTING_IMAGE = asset_image("lifting_plan.jpg")
RISK_IMAGE = asset_image("risk_assessment.jpg", "method_statement.jpg")
GEAR_IMAGE = asset_image("gear_register.jpg", "lifting_plan.jpg")
TRAINING_IMAGE = asset_image("training_certificate.jpg", "method_statement.jpg")
EXPIRY_IMAGE = asset_image("expiry_alert.jpg", "training_certificate.jpg")

# =====================
# PROFESSIONAL UI STYLE
# =====================
st.markdown(f"""
<style>
.stApp {{
    background: linear-gradient(180deg, #f4f7fb 0%, #eef2f7 100%);
}}

.block-container {{
    padding-top: 1.1rem;
    padding-bottom: 2rem;
    max-width: 1520px;
}}

[data-testid="stSidebar"] {{
    background: linear-gradient(180deg, #071126 0%, #0f172a 52%, #1e293b 100%);
    border-right: 1px solid rgba(255,255,255,0.08);
}}

[data-testid="stSidebar"] * {{
    color: white;
}}

.ewmt-header {{
    position: relative;
    overflow: hidden;
    min-height: 155px;
    background:
        linear-gradient(90deg, rgba(15,23,42,0.96), rgba(30,58,138,0.90), rgba(15,23,42,0.76)),
        url("data:image/jpg;base64,{HEADER_IMAGE}");
    background-size: cover;
    background-position: center;
    padding: 30px 36px;
    border-radius: 24px;
    color: white;
    margin-bottom: 24px;
    box-shadow: 0px 16px 38px rgba(15,23,42,0.26);
    border: 1px solid rgba(255,255,255,0.14);
}}

.ewmt-header:after {{
    content: "";
    position: absolute;
    left: 36px;
    right: 36px;
    bottom: 0;
    height: 4px;
    background: linear-gradient(90deg, #f59e0b, rgba(245,158,11,0.15));
    border-radius: 99px;
}}

.ewmt-badge {{
    display: inline-block;
    background: rgba(255,255,255,0.14);
    border: 1px solid rgba(255,255,255,0.22);
    padding: 7px 12px;
    border-radius: 999px;
    font-size: 13px;
    color: #dbeafe;
    margin-bottom: 10px;
    backdrop-filter: blur(6px);
}}

.ewmt-title {{
    font-size: 36px;
    line-height: 1.15;
    font-weight: 900;
    margin-bottom: 8px;
    letter-spacing: -0.6px;
}}

.ewmt-subtitle {{
    font-size: 17px;
    color: #dbeafe;
    max-width: 850px;
}}

.section-title {{
    font-size: 25px;
    font-weight: 850;
    color: #0f172a;
    margin-top: 18px;
    margin-bottom: 8px;
}}

.section-caption {{
    color: #64748b;
    margin-bottom: 18px;
}}

.metric-card {{
    background: rgba(255,255,255,0.92);
    border: 1px solid #e2e8f0;
    border-radius: 18px;
    padding: 18px 20px;
    box-shadow: 0px 8px 24px rgba(15,23,42,0.07);
    position: relative;
    overflow: hidden;
    min-height: 105px;
}}

.metric-card:before {{
    content: "";
    position: absolute;
    left: 0;
    top: 0;
    bottom: 0;
    width: 5px;
    background: #f59e0b;
}}

.metric-label {{
    color: #64748b;
    font-size: 14px;
    font-weight: 700;
}}

.metric-value {{
    color: #0f172a;
    font-size: 32px;
    font-weight: 900;
    margin-top: 6px;
}}

.metric-small {{
    color: #94a3b8;
    font-size: 12px;
    margin-top: 2px;
}}

.dashboard-card {{
    background: white;
    border-radius: 22px;
    border: 1px solid #e2e8f0;
    box-shadow: 0px 10px 28px rgba(15,23,42,0.08);
    overflow: hidden;
    margin-bottom: 12px;
    min-height: 330px;
    transition: all 0.18s ease;
}}

.dashboard-card:hover {{
    transform: translateY(-4px);
    box-shadow: 0px 18px 38px rgba(15,23,42,0.14);
    border-color: rgba(30,58,138,0.28);
}}

.dashboard-img {{
    height: 145px;
    background-size: cover;
    background-position: center;
    position: relative;
}}

.dashboard-img:after {{
    content: "";
    position: absolute;
    inset: 0;
    background: linear-gradient(180deg, rgba(15,23,42,0.05), rgba(15,23,42,0.74));
}}

.dashboard-pill {{
    position: absolute;
    left: 18px;
    bottom: 14px;
    z-index: 2;
    color: white;
    background: rgba(15,23,42,0.70);
    border: 1px solid rgba(255,255,255,0.20);
    padding: 6px 10px;
    border-radius: 999px;
    font-size: 12px;
    font-weight: 800;
}}

.dashboard-content {{
    padding: 20px 22px 10px 22px;
}}

.dashboard-content h3 {{
    margin: 0 0 10px 0;
    color: #0f172a;
    font-size: 23px;
    font-weight: 900;
}}

.dashboard-content p {{
    color: #475569;
    font-size: 15px;
    line-height: 1.55;
    min-height: 52px;
    margin-bottom: 0;
}}

.card-accent {{
    height: 4px;
    width: 58px;
    background: #f59e0b;
    border-radius: 999px;
    margin-bottom: 14px;
}}

.stButton > button {{
    background: linear-gradient(90deg, #1e3a8a, #2563eb);
    color: white;
    border-radius: 12px;
    padding: 0.70rem 1.25rem;
    font-weight: 800;
    border: none;
    width: 100%;
    box-shadow: 0px 8px 18px rgba(30,58,138,0.20);
}}

.stButton > button:hover {{
    background: linear-gradient(90deg, #0f172a, #1e3a8a);
    color: white;
}}

.stDownloadButton > button {{
    background: linear-gradient(90deg, #047857, #059669);
    color: white;
    border-radius: 12px;
    padding: 0.70rem 1.25rem;
    font-weight: 800;
    border: none;
}}

.stDownloadButton > button:hover {{
    background: #065f46;
    color: white;
}}

div[data-testid="stExpander"] {{
    border-radius: 15px;
    border: 1px solid #e2e8f0;
    box-shadow: 0px 4px 14px rgba(15,23,42,0.04);
    background: white;
}}

.footer-note {{
    margin-top: 20px;
    padding: 18px 22px;
    background: #0f172a;
    color: #cbd5e1;
    border-radius: 18px;
    border-left: 5px solid #f59e0b;
}}
</style>
""", unsafe_allow_html=True)

# =====================
# HEADER
# =====================
st.markdown("""
<div class="ewmt-header">
    <div class="ewmt-badge">EWMT INTERNAL AI SYSTEM</div>
    <div class="ewmt-title">Eric Wong Machinery Transportation Pte Ltd</div>
    <div class="ewmt-subtitle">
        Heavy Machinery Moving • Lifting • Transportation • AI Document Control System
    </div>
</div>
""", unsafe_allow_html=True)

# =====================
# COMMON FUNCTIONS
# =====================
def go_to_page(page_name):
    st.session_state.page = page_name
    st.rerun()


def replace_all(doc, data):
    def replace_in_paragraph(paragraph, replacements):
        if not paragraph.runs:
            return

        full_text = "".join(run.text for run in paragraph.runs)
        original_text = full_text

        for k, v in replacements.items():
            full_text = full_text.replace(k, str(v))

        if full_text != original_text:
            for run in paragraph.runs:
                run.text = ""
            paragraph.runs[0].text = full_text

    def replace_in_table(table, replacements):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)

                for nested_table in cell.tables:
                    replace_in_table(nested_table, replacements)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, data)

    for table in doc.tables:
        replace_in_table(table, data)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(paragraph, data)

        for table in section.header.tables:
            replace_in_table(table, data)

        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(paragraph, data)

        for table in section.footer.tables:
            replace_in_table(table, data)


def tick(value):
    return "☑" if value else "☐"


def safe_text(value):
    return "" if value is None else str(value)


def clean_ms_text(text):
    banned_phrases = [
        "The above equipment selection",
        "These safety controls",
        "The above sequence",
        "This sequence",
        "This method statement",
        "consistent with Eric Wong Machinery Transportation Pte Ltd",
        "company’s established method statement style",
        "company's established method statement style",
        "standard precautions repeatedly stated",
        "reflect the standard precautions",
        "follows the company",
        "previous method statements",
        "prior method statements",
    ]

    lines = str(text).splitlines()
    cleaned = []

    for line in lines:
        if not any(phrase.lower() in line.lower() for phrase in banned_phrases):
            cleaned.append(line)

    return "\n".join(cleaned).strip()


def set_ra_cell_text(cell, text):
    cell.text = ""
    p = cell.paragraphs[0]

    for i, line in enumerate(str(text).split("\n")):
        if i > 0:
            p.add_run().add_break()

        run = p.add_run(line)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(10)


def format_risk_assessment(doc):
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(10)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = "Times New Roman"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                        run.font.size = Pt(10)


def find_ra_table(doc):
    for table in doc.tables:
        full = " ".join(c.text for row in table.rows for c in row.cells)
        if "Hazard Identification" in full and "Risk Evaluation" in full and "Risk Control" in full:
            return table
    return None


def find_ra_column_header_row(table):
    for i, row in enumerate(table.rows):
        texts = [cell.text.strip() for cell in row.cells]
        if "Ref" in texts and "Work Activity" in texts and "Hazard" in texts:
            return i
    return None


def clear_rows_after_column_header(table):
    header_index = find_ra_column_header_row(table)

    if header_index is None:
        raise Exception("Cannot find RA column header row")

    while len(table.rows) > header_index + 1:
        row = table.rows[header_index + 1]
        row._element.getparent().remove(row._element)


def add_ra_row(table, values):
    row = table.add_row().cells
    for i, v in enumerate(values):
        if i < len(row):
            set_ra_cell_text(row[i], v)


def merge_same_work_activity_cells(table):
    header_index = find_ra_column_header_row(table)
    if header_index is None:
        return

    start_row = header_index + 1
    end_row = len(table.rows) - 1
    current_start = start_row

    while current_start <= end_row:
        activity = table.rows[current_start].cells[1].text.strip()
        current_end = current_start

        while (
            current_end + 1 <= end_row
            and table.rows[current_end + 1].cells[1].text.strip() == ""
        ):
            current_end += 1

        if activity and current_end > current_start:
            table.rows[current_start].cells[0].merge(table.rows[current_end].cells[0])
            table.rows[current_start].cells[1].merge(table.rows[current_end].cells[1])

        current_start = current_end + 1


def find_inventory_table(doc):
    for table in doc.tables:
        full = " ".join(c.text for row in table.rows for c in row.cells)
        if "Ref No." in full and "Location" in full and "Process" in full and "S/No." in full and "Work Activity" in full:
            return table
    return None


def fill_inventory_table(doc, activities_text, location, process):
    table = find_inventory_table(doc)

    if table is None:
        return

    activity_list = [
        a.strip()
        for a in activities_text.split("\n")
        if a.strip()
    ]

    start_row = None
    for i, row in enumerate(table.rows):
        row_text = " ".join(cell.text.strip() for cell in row.cells)
        if "S/No." in row_text and "Work Activity" in row_text:
            start_row = i + 1
            break

    if start_row is None:
        return

    for idx, activity in enumerate(activity_list, start=1):
        row_index = start_row + idx - 1

        if row_index >= len(table.rows):
            table.add_row()

        row = table.rows[row_index].cells

        if len(row) >= 6:
            set_ra_cell_text(row[0], "")
            set_ra_cell_text(row[1], location if idx == 1 else "")
            set_ra_cell_text(row[2], process if idx == 1 else "")
            set_ra_cell_text(row[3], str(idx))
            set_ra_cell_text(row[4], activity)
            set_ra_cell_text(row[5], "")


def certificate_browser(folder_name, title, info_text, search_label, search_placeholder, download_label):
    st.markdown(f"## {title}")
    st.info(info_text)

    cert_folder = os.path.join(BASE_DIR, folder_name)

    if not os.path.exists(cert_folder):
        st.error(f"Folder not found: {folder_name}")
        st.code(folder_name)
        st.info(f"Create this folder in your GitHub project and upload files inside.")
        return

    files = [
        f for f in os.listdir(cert_folder)
        if f.lower().endswith((".pdf", ".png", ".jpg", ".jpeg", ".docx"))
    ]

    files = sorted(files)

    if not files:
        st.warning(f"No files found inside {folder_name} folder.")
        return

    st.success(f"Found {len(files)} file(s).")

    search = st.text_input(
        search_label,
        "",
        placeholder=search_placeholder
    )

    filtered_files = files

    if search:
        search_words = search.lower().split()
        filtered_files = [
            f for f in files
            if all(word in f.lower() for word in search_words)
        ]

    if not filtered_files:
        st.warning("No matching file found.")
        return

    st.success(f"Found {len(filtered_files)} matching file(s).")

    selected_file = st.selectbox("Choose file", filtered_files)
    file_path = os.path.join(cert_folder, selected_file)

    st.write("Selected file:")
    st.code(selected_file)

    with open(file_path, "rb") as f:
        file_bytes = f.read()

    mime_type = "application/octet-stream"

    if selected_file.lower().endswith(".pdf"):
        mime_type = "application/pdf"
    elif selected_file.lower().endswith(".docx"):
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif selected_file.lower().endswith(".png"):
        mime_type = "image/png"
    elif selected_file.lower().endswith((".jpg", ".jpeg")):
        mime_type = "image/jpeg"

    st.download_button(
        download_label,
        file_bytes,
        file_name=selected_file,
        mime=mime_type
    )

    if selected_file.lower().endswith((".png", ".jpg", ".jpeg")):
        st.image(file_path, caption=selected_file)

    if selected_file.lower().endswith(".pdf"):
        st.markdown("### PDF Preview")
        base64_pdf = base64.b64encode(file_bytes).decode("utf-8")

        st.markdown(
            f"""
            <a href="data:application/pdf;base64,{base64_pdf}"
               target="_blank"
               style="
                   display:inline-block;
                   background:#1e3a8a;
                   color:white;
                   padding:12px 20px;
                   border-radius:10px;
                   text-decoration:none;
                   font-weight:700;
               ">
               Open PDF Preview in New Tab
            </a>
            """,
            unsafe_allow_html=True
        )

        st.info("Chrome blocks embedded PDF preview in Streamlit. Click the button above to preview before downloading.")

    if selected_file.lower().endswith(".docx"):
        st.info("Word document preview is not supported inside Streamlit. Please download the file to view.")

# =====================
# GITHUB DOCUMENT MANAGEMENT
# =====================
def get_secret(name, default=None):
    """Safely read a Streamlit secret without exposing it in the UI."""
    try:
        return st.secrets[name]
    except Exception:
        return default


def github_settings():
    return {
        "token": get_secret("GITHUB_TOKEN", ""),
        "repo": get_secret("GITHUB_REPO", ""),
        "branch": get_secret("GITHUB_BRANCH", "main"),
    }


def github_headers():
    cfg = github_settings()
    return {
        "Authorization": f"Bearer {cfg['token']}",
        "Accept": "application/vnd.github+json",
        "X-GitHub-Api-Version": "2022-11-28",
        "User-Agent": "EWMT-Streamlit-Document-Manager",
    }


def github_api_url(path=""):
    cfg = github_settings()
    repo = cfg["repo"].strip().strip("/")
    safe_path = quote(path.strip("/"), safe="/")
    base = f"https://api.github.com/repos/{repo}/contents"
    return f"{base}/{safe_path}" if safe_path else base


def validate_github_config():
    cfg = github_settings()
    missing = []

    if not get_secret("ADMIN_PASSWORD", ""):
        missing.append("ADMIN_PASSWORD")
    if not cfg["token"]:
        missing.append("GITHUB_TOKEN")
    if not cfg["repo"]:
        missing.append("GITHUB_REPO")

    return missing


def github_get_file(path):
    """Return GitHub file metadata or None when the file does not exist."""
    cfg = github_settings()
    response = requests.get(
        github_api_url(path),
        headers=github_headers(),
        params={"ref": cfg["branch"]},
        timeout=30,
    )

    if response.status_code == 404:
        return None

    response.raise_for_status()
    return response.json()


def github_list_folder(folder_path):
    cfg = github_settings()
    response = requests.get(
        github_api_url(folder_path),
        headers=github_headers(),
        params={"ref": cfg["branch"]},
        timeout=30,
    )

    if response.status_code == 404:
        return []

    response.raise_for_status()
    data = response.json()

    if isinstance(data, dict):
        data = [data]

    return sorted(
        [item for item in data if item.get("type") == "file"],
        key=lambda item: item.get("name", "").lower(),
    )


def github_upload_file(folder_path, filename, file_bytes, replace_existing=False):
    cfg = github_settings()
    filename = os.path.basename(filename).strip()

    if not filename:
        raise ValueError("Invalid file name.")

    target_path = f"{folder_path.strip('/')}/{filename}"
    existing = github_get_file(target_path)

    if existing and not replace_existing:
        raise FileExistsError(
            "A file with the same name already exists. Tick 'Replace existing file' to overwrite it."
        )

    payload = {
        "message": f"EWMT document upload: {target_path}",
        "content": base64.b64encode(file_bytes).decode("utf-8"),
        "branch": cfg["branch"],
    }

    if existing:
        payload["sha"] = existing["sha"]

    response = requests.put(
        github_api_url(target_path),
        headers=github_headers(),
        json=payload,
        timeout=60,
    )

    if response.status_code not in (200, 201):
        try:
            detail = response.json().get("message", response.text)
        except Exception:
            detail = response.text
        raise RuntimeError(f"GitHub upload failed: {detail}")

    return response.json()


def github_delete_file(folder_path, filename, sha):
    cfg = github_settings()
    filename = os.path.basename(filename).strip()
    target_path = f"{folder_path.strip('/')}/{filename}"

    payload = {
        "message": f"EWMT document delete: {target_path}",
        "sha": sha,
        "branch": cfg["branch"],
    }

    response = requests.delete(
        github_api_url(target_path),
        headers=github_headers(),
        json=payload,
        timeout=60,
    )

    if response.status_code != 200:
        try:
            detail = response.json().get("message", response.text)
        except Exception:
            detail = response.text
        raise RuntimeError(f"GitHub delete failed: {detail}")

    return response.json()


def admin_is_logged_in():
    return bool(st.session_state.get("admin_authenticated", False))


def admin_login_form():
    st.markdown("### 🔐 Administrator Access")
    st.caption("Only authorised administrators can upload, replace or delete company documents.")

    with st.form("admin_login_form", clear_on_submit=False):
        password = st.text_input("Administrator Password", type="password")
        login = st.form_submit_button("🔓 Login as Administrator", use_container_width=True)

    if login:
        saved_password = str(get_secret("ADMIN_PASSWORD", ""))

        if not saved_password:
            st.error("ADMIN_PASSWORD is not configured in Streamlit Secrets.")
            return False

        if hmac.compare_digest(str(password), saved_password):
            st.session_state.admin_authenticated = True
            st.success("Administrator access granted.")
            st.rerun()
        else:
            st.error("Incorrect administrator password.")

    return False


def render_admin_document_manager():
    st.markdown("### 📁 Admin Document Manager")
    st.caption(
        "Upload new certificates, replace renewed certificates and delete obsolete files. "
        "Changes are committed directly to your GitHub repository."
    )

    missing = validate_github_config()
    if missing:
        st.error("Missing Streamlit Secrets: " + ", ".join(missing))
        st.info(
            "Add ADMIN_PASSWORD, GITHUB_TOKEN, GITHUB_REPO and optionally GITHUB_BRANCH "
            "in Streamlit Cloud → App settings → Secrets."
        )
        return

    cfg = github_settings()

    top_left, top_right = st.columns([3, 1])
    with top_left:
        st.info(f"Repository: {cfg['repo']}  •  Branch: {cfg['branch']}")
    with top_right:
        if st.button("🔒 Logout", key="admin_logout", use_container_width=True):
            st.session_state.admin_authenticated = False
            st.rerun()

    folder_map = {
        "🧰 Lifting Gear Certificates": "Lifting Gears Certificate",
        "👷 Worker Training Certificates": "Workers Certificate",
    }

    selected_type = st.selectbox(
        "Document Category",
        list(folder_map.keys()),
        key="admin_document_category",
    )
    selected_folder = folder_map[selected_type]

    st.markdown(f"#### Current Files — `{selected_folder}`")

    try:
        github_files = github_list_folder(selected_folder)
    except Exception as e:
        st.error("Unable to read documents from GitHub.")
        st.exception(e)
        github_files = []

    if github_files:
        search_admin = st.text_input(
            "Search current files",
            placeholder="Type worker name, sling, shackle, expiry date...",
            key=f"admin_search_{selected_folder}",
        )

        shown_files = github_files
        if search_admin:
            words = search_admin.lower().split()
            shown_files = [
                item for item in github_files
                if all(word in item.get("name", "").lower() for word in words)
            ]

        st.caption(f"{len(shown_files)} file(s) shown • {len(github_files)} total")

        if not shown_files:
            st.warning("No matching file found.")
        else:
            for item in shown_files:
                filename = item.get("name", "")
                size_bytes = int(item.get("size", 0) or 0)
                size_kb = size_bytes / 1024

                with st.container(border=True):
                    c1, c2 = st.columns([5, 1])
                    with c1:
                        st.markdown(f"**{filename}**")
                        st.caption(f"{size_kb:,.1f} KB")
                    with c2:
                        if item.get("html_url"):
                            st.link_button(
                                "View",
                                item["html_url"],
                                use_container_width=True,
                            )

                    delete_confirm = st.checkbox(
                        f"I confirm I want to permanently delete {filename}",
                        key=f"delete_confirm_{selected_folder}_{item.get('sha', filename)}",
                    )

                    if st.button(
                        "🗑 Delete Permanently",
                        key=f"delete_{selected_folder}_{item.get('sha', filename)}",
                        disabled=not delete_confirm,
                        use_container_width=True,
                    ):
                        try:
                            with st.spinner(f"Deleting {filename}..."):
                                github_delete_file(selected_folder, filename, item["sha"])
                            st.success(f"Deleted: {filename}")
                            st.cache_data.clear()
                            st.rerun()
                        except Exception as e:
                            st.error("Delete failed.")
                            st.exception(e)
    else:
        st.warning("No documents found in this GitHub folder, or the folder does not exist yet.")

    st.markdown("---")
    st.markdown("#### ⬆ Upload / Replace Document")

    allowed_types = ["pdf", "png", "jpg", "jpeg", "docx"]
    uploaded_file = st.file_uploader(
        "Choose document",
        type=allowed_types,
        key=f"admin_upload_{selected_folder}",
    )

    if uploaded_file is not None:
        default_name = os.path.basename(uploaded_file.name)
        upload_name = st.text_input(
            "File name in GitHub",
            value=default_name,
            key=f"admin_filename_{selected_folder}",
        )

        replace_existing = st.checkbox(
            "Replace existing file if the same filename already exists",
            value=False,
            key=f"admin_replace_{selected_folder}",
        )

        st.info(
            "Tip: for lifting gear expiry alerts, include the expiry date in the filename, "
            "for example: 10 Ton Shackle Expiry 2027-08-31.pdf"
        )

        if st.button(
            "⬆ Upload Document",
            key=f"admin_upload_button_{selected_folder}",
            use_container_width=True,
        ):
            clean_name = os.path.basename(upload_name).strip()
            extension = clean_name.rsplit(".", 1)[-1].lower() if "." in clean_name else ""

            if not clean_name:
                st.error("Please enter a valid file name.")
            elif extension not in allowed_types:
                st.error("Allowed file types: PDF, PNG, JPG, JPEG and DOCX.")
            else:
                try:
                    with st.spinner(f"Uploading {clean_name}..."):
                        github_upload_file(
                            selected_folder,
                            clean_name,
                            uploaded_file.getvalue(),
                            replace_existing=replace_existing,
                        )
                    st.success(f"Uploaded successfully: {clean_name}")
                    st.info(
                        "GitHub has been updated. Streamlit Cloud normally redeploys automatically "
                        "after the repository changes."
                    )
                    st.cache_data.clear()
                except FileExistsError as e:
                    st.warning(str(e))
                except Exception as e:
                    st.error("Upload failed.")
                    st.exception(e)


# ======================================================
# PROFESSIONAL METHOD STATEMENT HELPERS
# ======================================================
MOS_BLUE = "0B5A7A"
MOS_DARK = "17324D"
MOS_LIGHT = "EAF2F7"
MOS_GREY = "EEF1F4"
MOS_RED = "C0392B"
MOS_GREEN = "2E7D32"


def mos_set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def mos_set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def mos_set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def mos_set_font(run, size=9, bold=False, color=None, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def mos_style_paragraph(paragraph, size=9, bold=False, color=None, align=None, space_after=4):
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.08
    for run in paragraph.runs:
        mos_set_font(run, size=size, bold=bold, color=color)


def mos_add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    mos_set_font(run, 8, color="64748B")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)

    run2 = paragraph.add_run(" of ")
    mos_set_font(run2, 8, color="64748B")
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "begin")
    instr_text2 = OxmlElement("w:instrText")
    instr_text2.set(qn("xml:space"), "preserve")
    instr_text2.text = " NUMPAGES "
    fld_char4 = OxmlElement("w:fldChar")
    fld_char4.set(qn("w:fldCharType"), "end")
    run2._r.append(fld_char3)
    run2._r.append(instr_text2)
    run2._r.append(fld_char4)


def mos_setup_document(doc, project):
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(9)

    for section in doc.sections:
        section.top_margin = Inches(0.88)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.62)
        section.right_margin = Inches(0.62)
        section.header_distance = Inches(0.18)
        section.footer_distance = Inches(0.25)

        header = section.header
        header.is_linked_to_previous = False
        if header.paragraphs:
            p0 = header.paragraphs[0]
            p0.text = ""
        table = header.add_table(rows=1, cols=2, width=Inches(7.0))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.columns[0].width = Inches(4.8)
        table.columns[1].width = Inches(2.2)
        left, right = table.rows[0].cells
        mos_set_cell_shading(left, "FFFFFF")
        mos_set_cell_shading(right, "FFFFFF")

        logo_candidates = [
            os.path.join(ASSET_DIR, "logo.png"),
            os.path.join(ASSET_DIR, "logo.jpg"),
            os.path.join(ASSET_DIR, "ewmt_logo.png"),
        ]
        logo = next((p for p in logo_candidates if os.path.exists(p)), None)
        if logo:
            lp = left.paragraphs[0]
            lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            lp.add_run().add_picture(logo, width=Inches(0.55))
            p = left.add_paragraph()
        else:
            p = left.paragraphs[0]

        r = p.add_run("ERIC WONG MACHINERY TRANSPORTATION PTE LTD")
        mos_set_font(r, 10, True, MOS_DARK)
        p2 = left.add_paragraph("28 Kranji Loop #05-03, Kranji Green, Singapore 739571\nTel: 62824175, 62828203, 62826020  •  www.ericwong.com")
        mos_style_paragraph(p2, size=7, color="475569", space_after=0)

        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr = rp.add_run("METHOD STATEMENT\n")
        mos_set_font(rr, 9, True, MOS_BLUE)
        rr2 = rp.add_run("LIFTING & MACHINERY MOVING")
        mos_set_font(rr2, 7, True, MOS_DARK)

        meta = header.add_table(rows=2, cols=4, width=Inches(7.0))
        meta.alignment = WD_TABLE_ALIGNMENT.CENTER
        labels = [
            ("Customer", project.get("customer", "")),
            ("Site", project.get("location", "")),
            ("Prepared", project.get("prepared_by", "")),
            ("Approved", project.get("approved_by", "")),
        ]
        for i, (lab, val) in enumerate(labels):
            row = 0 if i < 2 else 1
            col = (i % 2) * 2
            c1, c2 = meta.rows[row].cells[col], meta.rows[row].cells[col + 1]
            mos_set_cell_shading(c1, MOS_BLUE)
            mos_set_cell_shading(c2, MOS_LIGHT)
            c1.text = lab
            c2.text = str(val)
            for c in (c1, c2):
                mos_set_cell_margins(c, top=45, bottom=45, start=70, end=70)
                c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in c1.paragraphs[0].runs:
                mos_set_font(run, 6.5, True, "FFFFFF")
            for run in c2.paragraphs[0].runs:
                mos_set_font(run, 6.5, False, MOS_DARK)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = "EWMT Internal Controlled Document  •  "
        mos_style_paragraph(fp, size=7, color="64748B", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
        mos_add_page_field(fp)


def mos_add_title(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    if level == 1:
        r = p.add_run(text)
        mos_set_font(r, 15, True, MOS_BLUE)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(7)
    else:
        r = p.add_run(text)
        mos_set_font(r, 11, True, MOS_DARK)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
    return p


def mos_add_body(doc, text, bullet=False, bold_prefix=None):
    if not str(text).strip():
        return
    p = doc.add_paragraph(style=None)
    if bullet:
        p.style = doc.styles["List Bullet"]
    r = p.add_run(str(text).strip())
    mos_set_font(r, 9, False, "1F2937")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    return p


def mos_add_bullets(doc, items):
    for item in items or []:
        if str(item).strip():
            mos_add_body(doc, str(item).strip(), bullet=True)


def mos_add_key_value_table(doc, rows, widths=(2.2, 4.8)):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(widths[0])
        cells[1].width = Inches(widths[1])
        mos_set_cell_shading(cells[0], MOS_LIGHT)
        cells[0].text = str(label)
        cells[1].text = str(value if value not in (None, "") else "-")
        for run in cells[0].paragraphs[0].runs:
            mos_set_font(run, 8, True, MOS_DARK)
        for run in cells[1].paragraphs[0].runs:
            mos_set_font(run, 8, False, "1F2937")
        for c in cells:
            mos_set_cell_margins(c)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def mos_add_three_col_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        mos_set_cell_shading(cell, MOS_BLUE)
        for run in cell.paragraphs[0].runs:
            mos_set_font(run, 8, True, "FFFFFF")
    mos_set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for run in cells[i].paragraphs[0].runs:
                mos_set_font(run, 8, False, "1F2937")
            mos_set_cell_margins(cells[i])
    return table


def mos_get_font(size=28, bold=False):
    if ImageFont is None:
        return None
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    ]
    for candidate in candidates:
        try:
            if os.path.exists(candidate):
                return ImageFont.truetype(candidate, size=size)
        except Exception:
            pass
    return ImageFont.load_default()


def mos_arrow(draw, start, end, fill="#0B5A7A", width=8, head=24):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    angle = math.atan2(y2 - y1, x2 - x1)
    left = (x2 - head * math.cos(angle - math.pi / 6), y2 - head * math.sin(angle - math.pi / 6))
    right = (x2 - head * math.cos(angle + math.pi / 6), y2 - head * math.sin(angle + math.pi / 6))
    draw.polygon([end, left, right], fill=fill)


def mos_canvas(title):
    if Image is None:
        return None, None
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)
    title_font = mos_get_font(42, True)
    draw.text((60, 40), title, fill="#17324D", font=title_font)
    draw.line((60, 105, 1540, 105), fill="#0B5A7A", width=5)
    return img, draw


def mos_make_top_view(data):
    """Generate a simple top-view schematic without inventing a building when none is relevant."""
    mode = str(data.get("site_arrangement", ""))
    is_elevated = "Elevated" in mode
    has_opening = "Opening" in mode or "Door" in mode or "Roller Shutter" in mode

    title = "Top View – Lifting / Landing Arrangement"
    img, draw = mos_canvas(title)
    if img is None:
        return None

    f = mos_get_font(28, False)
    fb = mos_get_font(28, True)
    fs = mos_get_font(22, False)

    # Crane body and outriggers on the left.
    draw.rounded_rectangle((340, 330, 650, 600), radius=18, fill="#F4B942", outline="#7A4B00", width=5)
    draw.text((425, 450), "CRANE", fill="#4A2B00", font=fb)
    for y in (365, 565):
        draw.line((220, y, 790, y), fill="#444444", width=12)
        draw.ellipse((195, y-22, 240, y+22), fill="#B0B7BE", outline="#333333")
        draw.ellipse((775, y-22, 820, y+22), fill="#B0B7BE", outline="#333333")

    if is_elevated or has_opening:
        # A structure is only shown when the user says a building/opening is relevant.
        draw.rounded_rectangle((1030, 210, 1510, 720), radius=16, fill="#D7E7F0", outline="#17324D", width=5)
        draw.text((1155, 430), "STRUCTURE", fill="#17324D", font=fb)
        if has_opening:
            draw.rectangle((1020, 385, 1050, 545), fill="white", outline="#C0392B", width=4)
            if data.get('door_width', 0) > 0:
                draw.text((1060, 350), f"Opening {data.get('door_width', 0):.1f} m W", fill="#C0392B", font=f)

        draw.line((590, 420, 1130, 465), fill="#C0392B", width=18)
        draw.ellipse((1110, 445, 1150, 485), fill="#111111")
        draw.rounded_rectangle((1170, 420, 1385, 520), radius=10, fill="#E8EEF2", outline="#111111", width=4)
        draw.text((1210, 450), "LOAD", fill="#111111", font=fb)

        if data.get('crane_distance', 0) > 0:
            draw.text((820, 665), f"Crane to structure ≈ {data.get('crane_distance', 0):.1f} m", fill="#17324D", font=f)
        if data.get('landing_depth', 0) > 0:
            draw.text((1030, 755), f"Landing depth ≈ {data.get('landing_depth', 0):.1f} m", fill="#17324D", font=f)
    else:
        # Ground-level job: show pick-up and landing only, no artificial building.
        draw.rounded_rectangle((920, 285, 1160, 420), radius=14, fill="#E8EEF2", outline="#17324D", width=5)
        draw.text((965, 335), "PICK-UP\nLOAD", fill="#17324D", font=fb, spacing=4)
        draw.rounded_rectangle((1230, 510, 1510, 665), radius=14, fill="#DDEFE0", outline="#2E7D32", width=5)
        final_label = data.get("route_final") or "GROUND LANDING"
        draw.multiline_text((1260, 555), textwrap.fill(final_label, 18), fill="#1B5E20", font=f, spacing=4)
        draw.line((590, 420, 980, 350), fill="#C0392B", width=18)
        mos_arrow(draw, (1045, 430), (1290, 540), fill="#0B5A7A", width=8, head=24)
        draw.rounded_rectangle((865, 225, 1540, 710), radius=20, outline="#C0392B", width=4)
        draw.text((950, 725), "BARRICADED / EXCLUSION ZONE", fill="#C0392B", font=fb)

    if data.get('operating_radius', 0) > 0:
        draw.text((690, 790), f"Operating radius entered: {data.get('operating_radius', 0):.1f} m", fill="#17324D", font=f)

    draw.text((70, 835), "Planning schematic only – verify actual site dimensions, crane set-up, load chart and clearances before work.", fill="#64748B", font=fs)
    result = BytesIO()
    img.save(result, format="PNG")
    result.seek(0)
    return result


def mos_make_side_view(data):
    """Generate elevated geometry only when elevated work was selected; otherwise show a ground-level lift."""
    mode = str(data.get("site_arrangement", ""))
    is_elevated = "Elevated" in mode
    has_opening = "Opening" in mode or "Door" in mode or "Roller Shutter" in mode

    img, draw = mos_canvas("Side Elevation – Lifting Geometry" if is_elevated else "Side Elevation – Ground-Level Hoisting")
    if img is None:
        return None

    f = mos_get_font(28, False)
    fb = mos_get_font(28, True)
    fs = mos_get_font(22, False)
    ground_y = 745
    draw.line((70, ground_y, 1530, ground_y), fill="#555555", width=5)

    draw.rounded_rectangle((260, 590, 550, ground_y), radius=12, fill="#F4B942", outline="#7A4B00", width=5)
    draw.text((330, 645), "CRANE", fill="#4A2B00", font=fb)
    pivot = (490, 610)

    if is_elevated:
        draw.rectangle((1110, 180, 1500, ground_y), fill="#D7E7F0", outline="#17324D", width=5)
        draw.text((1220, 430), "STRUCTURE", fill="#17324D", font=fb)
        landing_y = 300
        if has_opening:
            draw.rectangle((1100, landing_y-65, 1130, landing_y+65), fill="white", outline="#C0392B", width=4)
        elbow = (780, 265)
        tip = (1130, landing_y)
        draw.line([pivot, elbow, tip], fill="#C0392B", width=20, joint="curve")
        draw.rectangle((1135, landing_y-40, 1340, landing_y+40), fill="#E8EEF2", outline="#111111", width=4)
        draw.text((1190, landing_y-18), "LOAD", fill="#111111", font=f)

        if data.get('building_height', 0) > 0:
            mos_arrow(draw, (1520, ground_y), (1520, 185), fill="#0B5A7A", width=5, head=18)
            draw.text((1280, 130), f"Lift height {data.get('building_height',0):.1f} m", fill="#0B5A7A", font=f)
        if data.get('crane_distance', 0) > 0:
            mos_arrow(draw, (550, 800), (1110, 800), fill="#0B5A7A", width=5, head=18)
            draw.text((650, 815), f"Crane to structure {data.get('crane_distance',0):.1f} m", fill="#0B5A7A", font=f)
    else:
        # Ground-level: no building height is shown.
        load_y = 600
        draw.rounded_rectangle((1010, load_y-65, 1250, load_y+65), radius=12, fill="#E8EEF2", outline="#17324D", width=5)
        draw.text((1070, load_y-15), "LOAD", fill="#17324D", font=fb)
        elbow = (760, 330)
        tip = (1130, load_y-65)
        draw.line([pivot, elbow, tip], fill="#C0392B", width=20, joint="curve")
        mos_arrow(draw, (1250, load_y), (1430, load_y), fill="#0B5A7A", width=7, head=22)
        draw.text((1270, load_y-50), "GROUND LANDING", fill="#2E7D32", font=f)

    draw.text((80, 835), f"Boom / jib length entered: {data.get('boom_length',0):.1f} m  •  Operating radius: {data.get('operating_radius',0):.1f} m", fill="#64748B", font=fs)
    result = BytesIO()
    img.save(result, format="PNG")
    result.seek(0)
    return result


def mos_make_route_view(data):
    img, draw = mos_canvas("Machinery Movement Route – Schematic")
    if img is None:
        return None

    f = mos_get_font(27, False)
    mode = str(data.get("site_arrangement", ""))
    has_opening = "Opening" in mode or "Door" in mode or "Roller Shutter" in mode
    points = [(180, 600), (510, 600), (770, 430), (1050, 430), (1350, 260)]
    labels = [
        data.get("route_start", "Starting Position") or "Starting Position",
        "Door / Opening" if has_opening else "Transfer Point",
        "Movement Route",
        "Turning / Alignment",
        data.get("route_final", "Final Position") or "Final Position",
    ]

    for i, point in enumerate(points):
        x, y = point
        draw.rounded_rectangle((x-105, y-45, x+105, y+45), radius=16, fill="#EAF2F7", outline="#17324D", width=4)
        draw.multiline_text((x-90, y-25), textwrap.fill(labels[i], width=18), fill="#17324D", font=f, align="center", spacing=4)
        if i < len(points)-1:
            mos_arrow(draw, (x+110, y), (points[i+1][0]-110, points[i+1][1]), fill="#0B5A7A", width=8)

    if data.get("route_notes"):
        draw.rounded_rectangle((140, 705, 1450, 830), radius=15, fill="#F8FAFC", outline="#94A3B8", width=3)
        draw.multiline_text((175, 730), textwrap.fill(data.get("route_notes"), 105), fill="#334155", font=mos_get_font(23), spacing=5)

    result = BytesIO()
    img.save(result, format="PNG")
    result.seek(0)
    return result


def mos_make_rigging_view(data):
    img, draw = mos_canvas("Rigging Arrangement – Schematic")
    if img is None:
        return None
    f = mos_get_font(28, False)
    fb = mos_get_font(28, True)
    hook = (800, 220)
    draw.ellipse((770, 175, 830, 235), outline="#111111", width=6)
    draw.line((800, 235, 800, 280), fill="#111111", width=8)
    load_box = (410, 570, 1190, 735)
    draw.rounded_rectangle(load_box, radius=12, fill="#E8EEF2", outline="#17324D", width=5)
    draw.text((700, 635), data.get("machine_name", "MACHINE / LOAD")[:24], fill="#17324D", font=fb)
    left_point = (525, 570)
    right_point = (1075, 570)
    draw.line((800, 280, left_point[0], left_point[1]), fill="#C0392B", width=12)
    draw.line((800, 280, right_point[0], right_point[1]), fill="#C0392B", width=12)
    draw.ellipse((left_point[0]-12, left_point[1]-12, left_point[0]+12, left_point[1]+12), fill="#111111")
    draw.ellipse((right_point[0]-12, right_point[1]-12, right_point[0]+12, right_point[1]+12), fill="#111111")
    # CG
    draw.line((800, 520, 800, 745), fill="#0B5A7A", width=4)
    draw.ellipse((780, 635, 820, 675), outline="#0B5A7A", width=5)
    draw.text((835, 635), "CG", fill="#0B5A7A", font=fb)
    draw.text((500, 330), f"Sling angle: {data.get('sling_angle',60):.0f}° (entered)", fill="#17324D", font=f)
    draw.text((500, 385), f"Lifting points: {data.get('lifting_points','Verify approved lifting points')}", fill="#17324D", font=f)
    draw.text((500, 440), f"Lifting gear: {str(data.get('lifting_gear',''))[:65]}", fill="#17324D", font=f)
    draw.text((70, 805), "Schematic only – lifting supervisor to verify actual rigging arrangement, sling angle, SWL and centre of gravity before lifting.", fill="#64748B", font=mos_get_font(22))
    out = BytesIO()
    img.save(out, format="PNG")
    out.seek(0)
    return out


def mos_add_picture_bytes(doc, image_bytes, caption=None, width=6.7):
    if not image_bytes:
        return
    try:
        image_bytes.seek(0)
    except Exception:
        pass
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(image_bytes, width=Inches(width))
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mos_style_paragraph(cp, size=8, color="64748B", space_after=7)


def mos_uploaded_image_bytes(uploaded_file):
    raw = uploaded_file.getvalue()
    if Image is None:
        return BytesIO(raw)
    try:
        im = Image.open(BytesIO(raw))
        if im.mode not in ("RGB", "L"):
            im = im.convert("RGB")
        out = BytesIO()
        im.save(out, format="JPEG", quality=88)
        out.seek(0)
        return out
    except Exception:
        return BytesIO(raw)


def mos_download_github_item(item):
    if not item:
        return None
    url = item.get("download_url")
    if not url:
        meta = github_get_file(item.get("path", ""))
        if meta and meta.get("content"):
            return base64.b64decode(meta["content"])
        url = (meta or {}).get("download_url")
    if not url:
        return None
    r = requests.get(url, headers=github_headers(), timeout=60)
    r.raise_for_status()
    return r.content


def mos_append_attachment(doc, title, filename, raw_bytes, full_pdf=True):
    doc.add_page_break()
    mos_add_title(doc, title, 1)
    mos_add_body(doc, f"Attachment: {filename}")
    ext = os.path.splitext(filename.lower())[1]

    if ext in (".png", ".jpg", ".jpeg"):
        try:
            mos_add_picture_bytes(doc, BytesIO(raw_bytes), caption=filename, width=6.5)
            return True
        except Exception as exc:
            mos_add_body(doc, f"Preview could not be inserted: {exc}")
            return False

    if ext == ".pdf":
        try:
            import fitz  # PyMuPDF - optional dependency
            pdf = fitz.open(stream=raw_bytes, filetype="pdf")
            max_pages = len(pdf) if full_pdf else min(1, len(pdf))
            for page_index in range(max_pages):
                page = pdf.load_page(page_index)
                pix = page.get_pixmap(matrix=fitz.Matrix(1.55, 1.55), alpha=False)
                png = BytesIO(pix.tobytes("png"))
                mos_add_picture_bytes(doc, png, caption=f"{filename} — page {page_index + 1} of {len(pdf)}", width=6.45)
                if page_index < max_pages - 1:
                    doc.add_page_break()
            pdf.close()
            return True
        except ImportError:
            mos_add_body(doc, "PDF preview embedding requires PyMuPDF in requirements.txt. The file name is included, but the PDF pages were not embedded.")
            return False
        except Exception as exc:
            mos_add_body(doc, f"PDF preview could not be inserted: {exc}")
            return False

    mos_add_body(doc, "This attachment type is listed but not embedded. Convert it to PDF/PNG/JPG if you want the full pages inside the MOS.")
    return False


def mos_site_prompt_text(data):
    """Create a clean site summary for AI; non-applicable dimensions are omitted rather than shown as zero."""
    lines = [
        f"Site arrangement: {data.get('site_arrangement', '')}",
        f"Ground / floor: {data.get('ground_condition', '')}",
        f"Access route: {data.get('access_route', '')}",
        f"Obstacles / interfaces: {data.get('obstacles', '')}",
        f"Environment: {data.get('environment', '')}",
    ]
    mode = str(data.get('site_arrangement', ''))
    is_elevated = 'Elevated' in mode
    has_opening = 'Opening' in mode or 'Door' in mode or 'Roller Shutter' in mode

    if is_elevated:
        if data.get('building_height', 0) > 0:
            lines.append(f"Lift / building height: {data.get('building_height')} m")
        if data.get('crane_distance', 0) > 0:
            lines.append(f"Crane-to-structure distance: {data.get('crane_distance')} m")
        if data.get('landing_depth', 0) > 0:
            lines.append(f"Landing depth: {data.get('landing_depth')} m")

    if has_opening and (data.get('door_width', 0) > 0 or data.get('door_height', 0) > 0):
        lines.append(f"Opening clearance: {data.get('door_width')} m W x {data.get('door_height')} m H")

    if data.get('route_start'):
        lines.append(f"Movement / pick-up start: {data.get('route_start')}")
    if data.get('route_final'):
        lines.append(f"Final / landing position: {data.get('route_final')}")
    if data.get('route_notes'):
        lines.append(f"Route / hoisting notes: {data.get('route_notes')}")

    return "\\n".join(lines)


def mos_site_condition_rows(data):
    """Create Word table rows only for site information that is relevant to this job."""
    rows = [
        ("Site / Hoisting Arrangement", data.get("site_arrangement")),
        ("Ground / Floor", data.get("ground_condition")),
        ("Access Route", data.get("access_route")),
        ("Obstacles / Interfaces", data.get("obstacles")),
        ("Environmental Controls", data.get("environment")),
    ]

    mode = str(data.get('site_arrangement', ''))
    is_elevated = 'Elevated' in mode
    has_opening = 'Opening' in mode or 'Door' in mode or 'Roller Shutter' in mode

    if is_elevated:
        if data.get('building_height', 0) > 0:
            rows.append(("Building / Lift Height", f"{data.get('building_height',0):.2f} m"))
        if data.get('crane_distance', 0) > 0:
            rows.append(("Crane-to-Structure Distance", f"{data.get('crane_distance',0):.2f} m"))
        if data.get('landing_depth', 0) > 0:
            rows.append(("Landing Depth", f"{data.get('landing_depth',0):.2f} m"))

    if has_opening and (data.get('door_width', 0) > 0 or data.get('door_height', 0) > 0):
        rows.append(("Opening Clearance", f"{data.get('door_width',0):.2f} m W × {data.get('door_height',0):.2f} m H"))

    route_parts = [str(x).strip() for x in [data.get('route_start'), data.get('route_final')] if str(x or '').strip()]
    if route_parts:
        rows.append(("Movement / Landing Arrangement", " → ".join(route_parts)))
    if data.get('route_notes'):
        rows.append(("Movement / Hoisting Notes", data.get('route_notes')))

    rows.append(("Lifting Crew", data.get("lifting_crew")))
    return rows


def mos_generate_ai_content(data):
    selected_equipment = ", ".join(data.get("equipment", []))
    site_details = mos_site_prompt_text(data)
    prompt = f"""
Prepare PROJECT-SPECIFIC professional Method Statement content for Eric Wong Machinery Transportation Pte Ltd.
Use Singapore machinery moving / lifting contractor wording.

PROJECT
Customer: {data.get('customer')}
Project: {data.get('project_name')}
Location: {data.get('location')}
Description: {data.get('description')}
Process: {data.get('process')}
Operation date/time: {data.get('operation_date')} {data.get('operation_time')}

LOAD
Machine / load: {data.get('machine_name')}
Load type: {data.get('load_type')}
Weight: {data.get('load_weight_kg')} kg
Rigging weight: {data.get('rigging_weight_kg')} kg
Dimensions L/W/H: {data.get('length_m')} / {data.get('width_m')} / {data.get('height_m')} m
Centre of gravity: {data.get('cg')}
Lifting points: {data.get('lifting_points')}
Lift classification: {data.get('lift_classification')}

SELECTED EQUIPMENT - ONLY USE THESE UNLESS A SAFETY ITEM IS NECESSARY
{selected_equipment}
Crane: {data.get('crane_model')}
Operating radius: {data.get('operating_radius')} m
SWL at operating radius: {data.get('swl_at_radius_kg')} kg
Boom length: {data.get('boom_length')} m
Lifting gear: {data.get('lifting_gear')}
Forklift: {data.get('forklift_details')}

SITE
{site_details}

PERSONNEL
{data.get('lifting_crew')}

Rules:
- Be site specific. Do not produce a generic equipment shopping list.
- Do not state "where required" repeatedly. Mention equipment that was actually selected.
- Work sequence should normally have 12 to 25 precise steps.
- Include trial lift where a crane/lorry loader lift is selected.
- Include floor protection / steel plates / plywood only if selected or stated in the site information.
- State hold points that the lifting supervisor must verify before proceeding.
- Include STOP WORK triggers for unsafe conditions.
- Do not claim engineering calculations have been performed unless the provided numbers show them.
- Do not invent certificates, registration numbers, load chart capacity, dimensions, personnel names or client requirements.
- If a site dimension is not provided because it is not applicable, do not mention it and do not invent a building, doorway or elevated landing.
- Return JSON only.
"""
    response = client.responses.create(
        model="gpt-5.4",
        input=prompt,
        tools=[{
            "type": "file_search",
            "vector_store_ids": [MS_VECTOR_STORE_ID]
        }],
        text={
            "format": {
                "type": "json_schema",
                "name": "professional_mos_schema",
                "schema": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "executive_summary": {"type": "string"},
                        "equipment_notes": {"type": "array", "items": {"type": "string"}},
                        "site_controls": {"type": "array", "items": {"type": "string"}},
                        "safety_controls": {"type": "array", "items": {"type": "string"}},
                        "work_sequence": {"type": "array", "items": {"type": "string"}},
                        "hold_points": {"type": "array", "items": {"type": "string"}},
                        "emergency_response": {"type": "array", "items": {"type": "string"}},
                        "assumptions": {"type": "array", "items": {"type": "string"}}
                    },
                    "required": [
                        "executive_summary", "equipment_notes", "site_controls", "safety_controls",
                        "work_sequence", "hold_points", "emergency_response", "assumptions"
                    ]
                }
            }
        }
    )
    return json.loads(response.output_text)


def mos_build_professional_doc(data, ai, site_photos, photo_captions, site_plan_file, cert_items, extra_attachments, full_pdf=True):
    doc = Document()
    project_meta = {
        "customer": data.get("customer", ""),
        "location": data.get("location", ""),
        "prepared_by": data.get("prepared_by", ""),
        "approved_by": data.get("approved_by", ""),
    }
    mos_setup_document(doc, project_meta)

    # COVER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    r = p.add_run("METHOD STATEMENT,\nLIFTING PLAN SUPPORT &\nSAFE WORK PROCEDURE")
    mos_set_font(r, 24, True, MOS_DARK)
    p2 = doc.add_paragraph("For Lifting and Machinery Moving Operation")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mos_style_paragraph(p2, size=14, bold=True, color=MOS_BLUE, space_after=30)

    mos_add_key_value_table(doc, [
        ("Customer / Tenant Company", data.get("customer")),
        ("Project", data.get("project_name")),
        ("Site Location", data.get("location")),
        ("Process", data.get("process")),
        ("Operation Date / Time", f"{data.get('operation_date')}  {data.get('operation_time')}"),
        ("Prepared By", data.get("prepared_by")),
        ("Approved By", data.get("approved_by")),
        ("Last / Next Review", f"{data.get('last_review')} / {data.get('next_review')}"),
        ("Lift Classification", data.get("lift_classification")),
    ])
    doc.add_paragraph()
    cp = doc.add_paragraph("CONTROLLED DOCUMENT – Site conditions and lifting parameters shall be verified before commencement of work.")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mos_style_paragraph(cp, size=9, bold=True, color=MOS_RED, space_after=0)
    doc.add_page_break()

    # CONTENTS
    mos_add_title(doc, "Contents", 1)
    contents = [
        "1.0 Workplace Safety & Health / Objectives",
        "2.0 Scope and Responsibilities",
        "3.0 Site Survey and Risk Assessment",
        "4.0 Lift Classification and Planning Controls",
        "5.0 Project / Load / Equipment Information",
        "6.0 Site Conditions and Controls",
        "7.0 Detailed Work Method and Hold Points",
        "8.0 Emergency / Stop Work Procedure",
        "9.0 Engineering & Site Drawings",
        "10.0 Site Photographs",
        "11.0 Certificates / Supporting Attachments",
        "12.0 Approval / Pre-Start Verification",
    ]
    for item in contents:
        mos_add_body(doc, item)
    doc.add_page_break()

    # STANDARD SECTIONS
    mos_add_title(doc, "1.0 Workplace Safety & Health / Objectives", 1)
    mos_add_body(doc, "The lifting and machinery moving operation shall be planned, supervised and carried out so far as reasonably practicable without risk to personnel, property or surrounding operations. Applicable workplace safety requirements, approved risk controls, lifting plans and permit-to-work requirements shall be implemented before work commences.")
    mos_add_bullets(doc, [
        "Prevent injury, dropped loads, uncontrolled machinery movement and property damage.",
        "Use trained and competent personnel for lifting supervision, rigging, signalling and equipment operation.",
        "Ensure lifting appliances, lifting gears and moving equipment are suitable, inspected and within valid certification where certification is required.",
        "Stop work when site conditions differ materially from the approved plan or become unsafe.",
    ])

    mos_add_title(doc, "2.0 Scope and Responsibilities", 1)
    mos_add_body(doc, "This Method Statement covers mobilisation, preparation, lifting/unloading, machinery shifting, positioning, housekeeping and demobilisation for the project described in this document.")
    mos_add_three_col_table(doc, ["Role", "Primary Responsibility", "Project Requirement"], [
        ("Project Manager / Coordinator", "Overall planning, client coordination and implementation of the approved work method.", "Confirm scope, access, resources and interfaces."),
        ("Lifting Supervisor", "Control the lifting operation and verify conditions, equipment, rigging and personnel before lifting.", "Authority to stop work."),
        ("Operator", "Operate the assigned crane / lorry loader / forklift within approved limits.", "Conduct pre-use checks."),
        ("Rigger / Signalman", "Rig the load, control the load and give clear agreed signals.", "Maintain exclusion zone and tag-line control."),
        ("Workers", "Follow briefing, RA, SWP and supervisor instructions.", "Use PPE and remain clear of suspended / moving loads."),
    ])

    mos_add_title(doc, "3.0 Site Survey and Risk Assessment", 1)
    mos_add_body(doc, "A competent person shall verify the load, access route, ground/floor conditions, headroom, door clearance, crane or vehicle set-up area, landing point and final machinery position. Hazards identified during the survey shall be reflected in the Risk Assessment and actual site controls.")
    mos_add_bullets(doc, [
        "Verify load weight, dimensions, centre of gravity and suitable lifting / support points.",
        "Confirm travel route capacity, turning space, slopes, floor protection and door / overhead clearances.",
        "Confirm crane set-up location, outrigger support, operating radius and load-chart capacity where lifting equipment is used.",
        "Review adjacent operations, public access, overhead obstructions, services and environmental conditions.",
    ])

    mos_add_title(doc, "4.0 Lift Classification and Planning Controls", 1)
    mos_add_key_value_table(doc, [
        ("Classification", data.get("lift_classification")),
        ("Classification Reason / Notes", data.get("lift_reason")),
        ("Crane / Lorry Loader Model", data.get("crane_model")),
        ("Operating Radius", f"{data.get('operating_radius',0):.2f} m"),
        ("SWL at Radius", f"{data.get('swl_at_radius_kg',0):,.0f} kg"),
        ("Total Lifted Weight", f"{data.get('total_lifted_kg',0):,.0f} kg"),
        ("Calculated Utilisation", f"{data.get('utilisation_pct',0):.1f}%" if data.get('swl_at_radius_kg',0) else "Not calculated"),
    ])
    if data.get("swl_at_radius_kg", 0) > 0:
        if data.get("utilisation_pct", 0) > 100:
            wp = doc.add_paragraph("STOP: Entered lifted weight exceeds entered SWL at operating radius. The lift shall not proceed on these values.")
            mos_style_paragraph(wp, size=10, bold=True, color=MOS_RED)
        elif data.get("utilisation_pct", 0) > 75:
            wp = doc.add_paragraph("ATTENTION: Entered utilisation is above 75%. Treat as elevated planning attention / non-routine where applicable to the project requirements.")
            mos_style_paragraph(wp, size=9, bold=True, color=MOS_RED)

    # PROJECT SPECIFIC
    doc.add_page_break()
    mos_add_title(doc, "5.0 Project / Load / Equipment Information", 1)
    mos_add_body(doc, ai.get("executive_summary", ""))
    mos_add_key_value_table(doc, [
        ("Description of Work", data.get("description")),
        ("Machine / Load", data.get("machine_name")),
        ("Load Type", data.get("load_type")),
        ("Load Weight", f"{data.get('load_weight_kg',0):,.0f} kg"),
        ("Rigging / Accessories Weight", f"{data.get('rigging_weight_kg',0):,.0f} kg"),
        ("Total Lifted Weight", f"{data.get('total_lifted_kg',0):,.0f} kg"),
        ("Dimensions (L × W × H)", f"{data.get('length_m',0):.2f} × {data.get('width_m',0):.2f} × {data.get('height_m',0):.2f} m"),
        ("Centre of Gravity", data.get("cg")),
        ("Lifting Points", data.get("lifting_points")),
    ])
    mos_add_title(doc, "Selected Equipment / Moving Tools", 2)
    mos_add_bullets(doc, data.get("equipment", []))
    mos_add_title(doc, "Lifting / Moving Equipment Details", 2)
    mos_add_key_value_table(doc, [
        ("Crane / Lorry Loader", data.get("crane_model")),
        ("LM / LE No.", data.get("crane_lm")),
        ("Maximum Capacity", data.get("crane_max_capacity")),
        ("Boom / Jib Length", f"{data.get('boom_length',0):.2f} m"),
        ("Operating Radius", f"{data.get('operating_radius',0):.2f} m"),
        ("SWL at Radius", f"{data.get('swl_at_radius_kg',0):,.0f} kg"),
        ("Forklift", data.get("forklift_details")),
        ("Lifting Gear", data.get("lifting_gear")),
    ])
    mos_add_bullets(doc, ai.get("equipment_notes", []))

    mos_add_title(doc, "6.0 Site Conditions and Controls", 1)
    mos_add_key_value_table(doc, mos_site_condition_rows(data))
    mos_add_bullets(doc, ai.get("site_controls", []))
    mos_add_title(doc, "Safety Controls", 2)
    mos_add_bullets(doc, ai.get("safety_controls", []))

    mos_add_title(doc, "7.0 Detailed Work Method and Hold Points", 1)
    for idx, step in enumerate(ai.get("work_sequence", []), start=1):
        mos_add_body(doc, f"{idx}. {step}")
    mos_add_title(doc, "Mandatory Hold Points", 2)
    mos_add_bullets(doc, ai.get("hold_points", []))

    mos_add_title(doc, "8.0 Emergency / Stop Work Procedure", 1)
    mos_add_bullets(doc, ai.get("emergency_response", []))
    if ai.get("assumptions"):
        mos_add_title(doc, "Planning Assumptions / Items to Verify", 2)
        mos_add_bullets(doc, ai.get("assumptions", []))

    # DRAWINGS
    doc.add_page_break()
    mos_add_title(doc, "9.0 Engineering & Site Drawings", 1)
    mos_add_body(doc, "The following drawings are planning schematics generated from the dimensions entered into the EWMT system. They are not Professional Engineer drawings. Actual dimensions, crane configuration, load chart capacity, rigging and site clearances shall be verified before work.")

    if site_plan_file is not None:
        mos_add_title(doc, "9.1 Uploaded Site Plan / Customer Drawing", 2)
        raw = site_plan_file.getvalue()
        ext = os.path.splitext(site_plan_file.name.lower())[1]
        if ext in (".png", ".jpg", ".jpeg"):
            mos_add_picture_bytes(doc, BytesIO(raw), caption=site_plan_file.name, width=6.5)
        elif ext == ".pdf":
            try:
                import fitz
                pdf = fitz.open(stream=raw, filetype="pdf")
                for pi in range(min(len(pdf), 3)):
                    pix = pdf.load_page(pi).get_pixmap(matrix=fitz.Matrix(1.4, 1.4), alpha=False)
                    mos_add_picture_bytes(doc, BytesIO(pix.tobytes("png")), caption=f"{site_plan_file.name} — page {pi+1}", width=6.4)
                pdf.close()
            except Exception as exc:
                mos_add_body(doc, f"Site-plan PDF preview unavailable: {exc}")

    drawing_specs = [
        ("9.2 Top View", data.get("draw_top"), mos_make_top_view, "Generated top-view schematic"),
        ("9.3 Side Elevation", data.get("draw_side"), mos_make_side_view, "Generated side-elevation schematic"),
        ("9.4 Machinery Movement Route", data.get("draw_route"), mos_make_route_view, "Generated machinery movement-route schematic"),
        ("9.5 Rigging Arrangement", data.get("draw_rigging"), mos_make_rigging_view, "Generated rigging schematic"),
    ]
    for title, enabled, fn, cap in drawing_specs:
        if enabled:
            mos_add_title(doc, title, 2)
            try:
                image_buf = fn(data)
                if image_buf:
                    mos_add_picture_bytes(doc, image_buf, caption=cap, width=6.6)
                else:
                    mos_add_body(doc, "Drawing library unavailable. Install Pillow to enable generated drawings.")
            except Exception as exc:
                mos_add_body(doc, f"Drawing could not be generated: {exc}")

    # PHOTOS
    doc.add_page_break()
    mos_add_title(doc, "10.0 Site Photographs", 1)
    if not site_photos:
        mos_add_body(doc, "No site photographs were uploaded for this issue.")
    else:
        for idx, photo in enumerate(site_photos, start=1):
            caption = photo_captions[idx-1] if idx-1 < len(photo_captions) and photo_captions[idx-1].strip() else f"Site Photograph {idx}: {photo.name}"
            mos_add_picture_bytes(doc, mos_uploaded_image_bytes(photo), caption=caption, width=6.35)

    # CERTS & ATTACHMENTS
    doc.add_page_break()
    mos_add_title(doc, "11.0 Certificates / Supporting Attachments", 1)
    if not cert_items and not extra_attachments:
        mos_add_body(doc, "No supporting certificates or attachments were selected for embedding.")
    else:
        mos_add_body(doc, "The following selected documents are appended for project reference. Validity and applicability shall be verified before deployment.")

    for category, item in cert_items:
        try:
            raw = mos_download_github_item(item)
            if raw:
                mos_append_attachment(doc, category, item.get("name", "certificate"), raw, full_pdf=full_pdf)
            else:
                mos_add_body(doc, f"Could not download selected file: {item.get('name','')}")
        except Exception as exc:
            mos_add_body(doc, f"Could not append {item.get('name','')}: {exc}")

    for uploaded in extra_attachments or []:
        try:
            mos_append_attachment(doc, "Supporting Attachment", uploaded.name, uploaded.getvalue(), full_pdf=full_pdf)
        except Exception as exc:
            mos_add_body(doc, f"Could not append {uploaded.name}: {exc}")

    # APPROVAL
    doc.add_page_break()
    mos_add_title(doc, "12.0 Approval / Pre-Start Verification", 1)
    mos_add_body(doc, "Before commencement, the lifting team shall review the Method Statement, Risk Assessment, lifting plan / permit-to-work where applicable, emergency arrangements and actual site conditions. Any material change shall be reviewed before proceeding.")
    mos_add_three_col_table(doc, ["Role", "Name", "Signature / Date"], [
        ("Prepared By", data.get("prepared_by", ""), ""),
        ("Lifting Supervisor", data.get("lifting_supervisor", ""), ""),
        ("Assessed / Project Manager", data.get("assessed_by", ""), ""),
        ("Approved By", data.get("approved_by", ""), ""),
    ])

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out


# =====================
# DASHBOARD COUNT FUNCTIONS
# =====================
def count_files_in_folder(folder_name, allowed_ext=(".pdf", ".png", ".jpg", ".jpeg", ".docx")):
    folder_path = os.path.join(BASE_DIR, folder_name)

    if not os.path.exists(folder_path):
        return 0

    return len([
        f for f in os.listdir(folder_path)
        if f.lower().endswith(allowed_ext)
    ])


def get_lifting_gear_expiry_counts(alert_days=30):
    import re

    cert_folder = os.path.join(BASE_DIR, "Lifting Gears Certificate")

    counts = {
        "expired": 0,
        "expiring_soon": 0,
        "valid": 0,
        "unknown": 0
    }

    if not os.path.exists(cert_folder):
        return counts

    files = [
        f for f in os.listdir(cert_folder)
        if f.lower().endswith((".pdf", ".png", ".jpg", ".jpeg"))
    ]

    today = date.today()

    patterns = [
        r"(\d{4})[-_\\.](\d{1,2})[-_\\.](\d{1,2})",
        r"(\d{1,2})[-_\\.](\d{1,2})[-_\\.](\d{4})",
    ]

    for f in files:
        found_date = None

        for pattern in patterns:
            match = re.search(pattern, f)

            if match:
                try:
                    parts = match.groups()

                    if len(parts[0]) == 4:
                        found_date = date(int(parts[0]), int(parts[1]), int(parts[2]))
                    else:
                        found_date = date(int(parts[2]), int(parts[1]), int(parts[0]))

                    break
                except Exception:
                    found_date = None

        if not found_date:
            counts["unknown"] += 1
        else:
            days_left = (found_date - today).days

            if days_left < 0:
                counts["expired"] += 1
            elif days_left <= alert_days:
                counts["expiring_soon"] += 1
            else:
                counts["valid"] += 1

    return counts

# =====================
# SIDEBAR NAVIGATION
# =====================
with st.sidebar:
    st.markdown("## EWMT System")
    st.markdown("AI Document Control")
    st.markdown("---")

    if st.button("🏠 Dashboard", key="side_dashboard"):
        go_to_page("🏠 Dashboard")

    if st.button("📄 Method Statement", key="side_method_statement"):
        go_to_page("📄 Method Statement")

    if st.button("🏗️ Lifting Plan", key="side_lifting_plan"):
        go_to_page("🏗️ Lifting Plan")

    if st.button("⚠️ Risk Assessment Pro", key="side_risk_assessment"):
        go_to_page("⚠️ Risk Assessment Pro")

    if st.button("🧰 Lifting Gear Register", key="side_lifting_gear"):
        go_to_page("🧰 Lifting Gear Register")

    if st.button("👷 Worker Training Certificate", key="side_worker_training"):
        go_to_page("👷 Worker Training Certificate")

    if st.button("⏰ Expiry Alerts", key="side_expiry_alerts"):
        go_to_page("⏰ Expiry Alerts")

    if st.button("⚙️ Settings", key="side_settings"):
        go_to_page("⚙️ Settings")

    page = st.session_state.page

    st.markdown("---")
    st.caption("Internal system for document preparation and lifting operation records.")
    
# ======================================================
# DASHBOARD
# ======================================================
if page == "🏠 Dashboard":
    lifting_gear_count = count_files_in_folder("Lifting Gears Certificate")
    worker_cert_count = count_files_in_folder("Workers Certificate")
    expiry_counts = get_lifting_gear_expiry_counts(alert_days=30)

    expired_count = expiry_counts["expired"]
    expiring_soon_count = expiry_counts["expiring_soon"]
    valid_count = expiry_counts["valid"]

    st.markdown("""
    <div class="section-title">EWMT AI Document Control Dashboard</div>
    <div class="section-caption">
        Professional document generation, lifting operation records and certificate control system.
    </div>
    """, unsafe_allow_html=True)

    m1, m2, m3, m4 = st.columns(4)

    with m1:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-label">Document Modules</div>
            <div class="metric-value">3</div>
            <div class="metric-small">Method Statement / RA / Lifting Plan</div>
        </div>
        """, unsafe_allow_html=True)

    with m2:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-label">Gear Records</div>
            <div class="metric-value">{lifting_gear_count}</div>
            <div class="metric-small">Files in Lifting Gears Certificate folder</div>
        </div>
        """, unsafe_allow_html=True)

    with m3:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-label">Expiring Soon</div>
            <div class="metric-value">{expiring_soon_count}</div>
            <div class="metric-small">Within next 30 days</div>
        </div>
        """, unsafe_allow_html=True)

    with m4:
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-label">Worker Certificates</div>
            <div class="metric-value">{worker_cert_count}</div>
            <div class="metric-small">Files in Workers Certificate folder</div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown('<div class="section-title">Document Modules</div>', unsafe_allow_html=True)

    def dashboard_card(title, desc, image_b64, tag):
        st.markdown(f"""
        <div class="dashboard-card">
            <div class="dashboard-img" style='background-image:url("data:image/jpg;base64,{image_b64}")'>
                <div class="dashboard-pill">{tag}</div>
            </div>
            <div class="dashboard-content">
                <div class="card-accent"></div>
                <h3>{title}</h3>
                <p>{desc}</p>
            </div>
        </div>
        """, unsafe_allow_html=True)

    col1, col2, col3 = st.columns(3)

    with col1:
        dashboard_card(
            "📄 Method Statement",
            "Create professional method statements for machinery moving, factory shifting, transport and lifting works.",
            METHOD_IMAGE,
            "WORK METHOD"
        )

        if st.button("Open Method Statement", key="open_ms"):
            go_to_page("📄 Method Statement")

    with col2:
        dashboard_card(
            "🏗️ Lifting Plan",
            "Generate lifting plan and permit-to-work documents based on load, crane, radius and site conditions.",
            LIFTING_IMAGE,
            "LIFTING OPERATION"
        )

        if st.button("Open Lifting Plan", key="open_lp"):
            go_to_page("🏗️ Lifting Plan")

    with col3:
        dashboard_card(
            "⚠️ Risk Assessment Pro",
            "Create structured 5x5 risk assessments using actual activity, hazard, controls and residual risk.",
            RISK_IMAGE,
            "SAFETY CONTROL"
        )

        if st.button("Open Risk Assessment", key="open_ra"):
            go_to_page("⚠️ Risk Assessment Pro")

    st.markdown('<div class="section-title">Certificate / Records Modules</div>', unsafe_allow_html=True)

    col4, col5, col6 = st.columns(3)

    with col4:
        dashboard_card(
            "🧰 Lifting Gear Register",
            "Manage shackles, slings, wire ropes, lifting certificates, SWL records and expiry dates.",
            GEAR_IMAGE,
            "GEAR RECORDS"
        )

        if st.button("Open Lifting Gear Register", key="open_lg"):
            go_to_page("🧰 Lifting Gear Register")

    with col5:
        dashboard_card(
            "👷 Worker Training Certificate",
            "Search, preview and download worker training certificates uploaded into your GitHub folders.",
            TRAINING_IMAGE,
            "WORKER RECORDS"
        )

        if st.button("Open Worker Training Certificate", key="open_worker_cert"):
            go_to_page("👷 Worker Training Certificate")

    with col6:
        dashboard_card(
            "⏰ Expiry Alerts",
            "Check expired and expiring lifting gear certificates using expiry dates in your file names.",
            EXPIRY_IMAGE,
            "EXPIRY MONITORING"
        )

        if st.button("Open Expiry Alerts", key="open_expiry"):
            go_to_page("⏰ Expiry Alerts")

    st.markdown('<div class="section-title">System</div>', unsafe_allow_html=True)

    col7, col8, col9 = st.columns(3)

    with col7:
        dashboard_card(
            "⚙️ Settings",
            "Manage template placeholders, prepared-by names and future default company details.",
            HEADER_IMAGE,
            "SYSTEM CONFIG"
        )

        if st.button("Open Settings", key="open_settings"):
            go_to_page("⚙️ Settings")

    st.markdown("""
    <div class="footer-note">
        <b>EWMT Internal System</b><br>
        Dashboard counts are now calculated from your GitHub folders and lifting gear expiry filenames.
    </div>
    """, unsafe_allow_html=True)
    
# ======================================================
# PROFESSIONAL METHOD STATEMENT
# ======================================================
if page == "📄 Method Statement":
    st.markdown("## 📄 Professional Method Statement Generator")
    st.caption(
        "Generate a complete project-specific MOS package with work method, site photos, "
        "automatic schematics and optional certificate appendices."
    )

    st.info(
        "V1 generates a professional Word package directly — no new Word template is required. "
        "Generated drawings are planning schematics and must be verified against the actual site / crane load chart."
    )

    with st.expander("1. Project & Document Control", expanded=True):
        c1, c2 = st.columns(2)
        with c1:
            ms_customer = st.text_input("Customer / Tenant Company", key="pms_customer")
            ms_project_name = st.text_input("Project Name", key="pms_project_name")
            ms_location = st.text_input("Site Location", key="pms_location")
            ms_description = st.text_area("Description of Work", height=110, key="pms_description")
            ms_process = st.text_input("Process", "Lifting and Moving of Machinery", key="pms_process")
        with c2:
            ms_date_input = st.date_input("Operation Date", value=date.today(), key="pms_date")
            ms_operation_time = st.text_input("Operation Time", key="pms_operation_time", placeholder="Example: 0900 hrs to 1700 hrs")
            ms_prepared_by = st.text_input("Prepared By", "Kevin Wong", key="pms_prepared_by")
            ms_assessed_by = st.text_input("Assessed / Project Manager", "Kevin Wong", key="pms_assessed_by")
            ms_approved_by = st.text_input("Approved By", "Eric Wong (Director)", key="pms_approved_by")
            ms_last_review = st.date_input("Last Review Date", value=date.today(), key="pms_last_review")
            ms_next_review = st.date_input("Next Review Date", value=date.today() + timedelta(days=730), key="pms_next_review")

    with st.expander("2. Load / Machine Details", expanded=True):
        c1, c2, c3 = st.columns(3)
        with c1:
            ms_machine_name = st.text_input("Machine / Load Name", key="pms_machine_name")
            ms_load_type = st.selectbox("Load Type", ["Machine", "Wooden Crate", "Machine Part", "Equipment", "Other"], key="pms_load_type")
            ms_load_weight = st.number_input("Load Weight (kg)", min_value=0.0, value=0.0, step=100.0, key="pms_load_weight")
            ms_rigging_weight = st.number_input("Rigging / Accessories Weight (kg)", min_value=0.0, value=0.0, step=50.0, key="pms_rigging_weight")
        with c2:
            ms_length = st.number_input("Length (m)", min_value=0.0, value=0.0, step=0.1, key="pms_length")
            ms_width = st.number_input("Width (m)", min_value=0.0, value=0.0, step=0.1, key="pms_width")
            ms_height = st.number_input("Height (m)", min_value=0.0, value=0.0, step=0.1, key="pms_height")
            ms_cg = st.text_input("Centre of Gravity", "Estimated at geometric centre unless manufacturer information states otherwise", key="pms_cg")
        with c3:
            ms_lifting_points = st.text_area("Lifting Points / Support Points", height=95, key="pms_lifting_points")
            ms_sling_angle = st.number_input("Sling Angle for Drawing (degrees)", min_value=30.0, max_value=90.0, value=60.0, step=5.0, key="pms_sling_angle")
            ms_lift_class = st.selectbox("Lift Classification", ["Routine Lift", "Non-Routine Lift", "To Be Confirmed"], key="pms_lift_class")
            ms_lift_reason = st.text_area("Classification Reason / Notes", height=75, key="pms_lift_reason")

        total_lifted = float(ms_load_weight) + float(ms_rigging_weight)
        st.metric("Total Lifted Weight", f"{total_lifted:,.0f} kg")

    with st.expander("3. Equipment Selection & Crane Data", expanded=True):
        st.markdown("**Select only equipment actually intended for this job.**")
        e1, e2, e3, e4 = st.columns(4)
        equipment = []
        with e1:
            if st.checkbox("Lorry Loader / Lorry Crane", key="pms_eq_lorry_crane"):
                equipment.append("Lorry loader / lorry crane")
            if st.checkbox("Mobile Crane", key="pms_eq_mobile_crane"):
                equipment.append("Mobile crane")
            if st.checkbox("Forklift", key="pms_eq_forklift"):
                equipment.append("Forklift")
        with e2:
            if st.checkbox("Hydraulic Jacks", value=True, key="pms_eq_jacks"):
                equipment.append("Hydraulic jacks")
            if st.checkbox("Machine Skates / Rollers", value=True, key="pms_eq_skates"):
                equipment.append("Machine skates / rollers")
            if st.checkbox("Pallet Truck", key="pms_eq_pallet"):
                equipment.append("Pallet truck")
        with e3:
            if st.checkbox("Chain Block / Lever Block", key="pms_eq_chain"):
                equipment.append("Chain block / lever block")
            if st.checkbox("Spreader Beam", key="pms_eq_spreader"):
                equipment.append("Spreader beam")
            if st.checkbox("Steel Plates", key="pms_eq_steel_plate"):
                equipment.append("Steel plates for floor / ground protection")
        with e4:
            if st.checkbox("Plywood / Timber Protection", value=True, key="pms_eq_plywood"):
                equipment.append("Plywood / timber floor protection")
            if st.checkbox("Barricades / Cones", value=True, key="pms_eq_barricade"):
                equipment.append("Barricades, warning signs and traffic cones")
            if st.checkbox("Tag Lines", value=True, key="pms_eq_tagline"):
                equipment.append("Tag lines")

        c1, c2, c3 = st.columns(3)
        with c1:
            ms_crane_model = st.text_input("Crane / Lorry Loader Make & Model", key="pms_crane_model")
            ms_crane_lm = st.text_input("LM / LE Registration No.", key="pms_crane_lm")
            ms_crane_max_capacity = st.text_input("Maximum Crane Capacity", key="pms_crane_max_capacity", placeholder="Example: 110 Ton")
        with c2:
            ms_boom_length = st.number_input("Boom / Jib Length (m)", min_value=0.0, value=0.0, step=0.5, key="pms_boom_length")
            ms_operating_radius = st.number_input("Operating Radius (m)", min_value=0.0, value=0.0, step=0.1, key="pms_radius")
            ms_swl_radius = st.number_input("SWL at Operating Radius (kg)", min_value=0.0, value=0.0, step=100.0, key="pms_swl_radius")
        with c3:
            ms_forklift_details = st.text_input("Forklift Details / Capacity", key="pms_forklift_details")
            ms_lifting_gear = st.text_area(
                "Lifting Gear / Moving Accessories",
                value="Certified webbing / wire rope slings, shackles and suitable rigging accessories",
                height=95,
                key="pms_lifting_gear"
            )

        utilisation = (total_lifted / ms_swl_radius * 100) if ms_swl_radius > 0 else 0.0
        m1, m2, m3 = st.columns(3)
        m1.metric("Total Lifted", f"{total_lifted:,.0f} kg")
        m2.metric("SWL @ Radius", f"{ms_swl_radius:,.0f} kg" if ms_swl_radius else "Not entered")
        m3.metric("Crane Utilisation", f"{utilisation:.1f}%" if ms_swl_radius else "Not calculated")
        if utilisation > 100:
            st.error("Entered total lifted weight exceeds the entered SWL at operating radius. Do not proceed on these figures.")
        elif utilisation > 75:
            st.warning("Entered utilisation is above 75%. Review lift classification and client/site requirements carefully.")

    with st.expander("4. Site Survey / Access / Drawing Details", expanded=True):
        # Keep the original free-text survey fields. Only the dimensions on the right are conditional.
        c1, c2 = st.columns(2)

        with c1:
            ms_ground = st.text_area(
                "Ground / Floor Condition",
                value="Firm, level and suitable for the intended equipment. Verify actual load-bearing condition before setup.",
                key="pms_ground"
            )
            ms_access = st.text_area(
                "Access Route",
                value="Confirm access route, turning space, headroom and final machine position before mobilisation.",
                key="pms_access"
            )
            ms_obstacles = st.text_area(
                "Obstacles / Interfaces",
                value="Clear obstructions in the working area and travel route. Barricade the operation area to prevent unauthorised entry.",
                key="pms_obstacles"
            )
            ms_environment = st.text_area(
                "Environment",
                value="Suspend lifting work during lightning, thunderstorms, heavy rain or any condition that makes the operation unsafe.",
                key="pms_environment"
            )

        with c2:
            ms_site_arrangement = st.selectbox(
                "Site / Hoisting Arrangement",
                [
                    "Ground / Floor Level - No Building Dimensions",
                    "Elevated / Upper Floor / Roof Hoisting",
                    "Ground Floor - Through Door / Roller Shutter / Opening",
                    "Elevated Hoist - Through Opening",
                    "Custom / Other",
                ],
                index=0,
                key="pms2_site_arrangement",
                help="Choose this only to decide which drawing dimensions are needed. It does not change the rest of your Method Statement form."
            )

            # Safe defaults when the dimensions are not relevant.
            ms_building_height = 0.0
            ms_crane_distance = 0.0
            ms_landing_depth = 0.0
            ms_door_width = 0.0
            ms_door_height = 0.0

            is_elevated_site = "Elevated" in ms_site_arrangement
            has_opening_site = "Opening" in ms_site_arrangement or "Door" in ms_site_arrangement or "Roller Shutter" in ms_site_arrangement

            if ms_site_arrangement == "Ground / Floor Level - No Building Dimensions":
                st.info("Ground-level job: building height, landing depth and door dimensions are not required.")

            elif is_elevated_site:
                st.markdown("**Elevated hoisting dimensions**")
                d1, d2 = st.columns(2)
                with d1:
                    ms_building_height = st.number_input(
                        "Lift / Building Height (m)", min_value=0.0, value=0.0, step=0.5, key="pms2_building_height"
                    )
                    ms_crane_distance = st.number_input(
                        "Crane to Structure Distance (m)", min_value=0.0, value=0.0, step=0.1, key="pms2_crane_distance"
                    )
                with d2:
                    ms_landing_depth = st.number_input(
                        "Landing Depth (m)", min_value=0.0, value=0.0, step=0.1, key="pms2_landing_depth"
                    )

                if has_opening_site:
                    o1, o2 = st.columns(2)
                    with o1:
                        ms_door_width = st.number_input(
                            "Opening Width (m)", min_value=0.0, value=0.0, step=0.1, key="pms2_door_width_elevated"
                        )
                    with o2:
                        ms_door_height = st.number_input(
                            "Opening Height (m)", min_value=0.0, value=0.0, step=0.1, key="pms2_door_height_elevated"
                        )

            elif has_opening_site:
                st.markdown("**Opening / access dimensions**")
                o1, o2 = st.columns(2)
                with o1:
                    ms_door_width = st.number_input(
                        "Opening Width (m)", min_value=0.0, value=0.0, step=0.1, key="pms2_door_width_ground"
                    )
                with o2:
                    ms_door_height = st.number_input(
                        "Opening Height (m)", min_value=0.0, value=0.0, step=0.1, key="pms2_door_height_ground"
                    )

            else:
                # Custom: show optional switches rather than forcing every dimension.
                need_elevated = st.checkbox("Need elevated / building dimensions", value=False, key="pms2_custom_elevated")
                need_opening = st.checkbox("Need door / opening dimensions", value=False, key="pms2_custom_opening")
                if need_elevated:
                    d1, d2, d3 = st.columns(3)
                    with d1:
                        ms_building_height = st.number_input("Lift Height (m)", min_value=0.0, value=0.0, step=0.5, key="pms2_custom_height")
                    with d2:
                        ms_crane_distance = st.number_input("Crane Distance (m)", min_value=0.0, value=0.0, step=0.1, key="pms2_custom_crane_distance")
                    with d3:
                        ms_landing_depth = st.number_input("Landing Depth (m)", min_value=0.0, value=0.0, step=0.1, key="pms2_custom_landing_depth")
                if need_opening:
                    o1, o2 = st.columns(2)
                    with o1:
                        ms_door_width = st.number_input("Opening Width (m)", min_value=0.0, value=0.0, step=0.1, key="pms2_custom_door_width")
                    with o2:
                        ms_door_height = st.number_input("Opening Height (m)", min_value=0.0, value=0.0, step=0.1, key="pms2_custom_door_height")

            ms_route_start = st.text_input(
                "Movement / Pick-Up Start Position",
                "Unloading / landing point",
                key="pms2_route_start"
            )
            ms_route_final = st.text_input(
                "Final / Landing Position",
                "Designated installation position",
                key="pms2_route_final"
            )
            ms_route_notes = st.text_area(
                "Movement / Hoisting Notes (optional)",
                height=90,
                key="pms2_route_notes",
                placeholder="Example: hoist from lorry and land beside factory entrance; or enter through Roller Shutter 2 and position beside Line 3."
            )

    with st.expander("5. Lifting Team", expanded=True):
        c1, c2 = st.columns(2)
        with c1:
            ms_lifting_supervisor = st.text_input("Lifting Supervisor", "Ibrahim / Zahari / Zaharin / Wong Yen Siong", key="pms_lifting_supervisor")
            ms_operator = st.text_input("Equipment Operator", "Lim Poh Soon / Norhalim / Lim Poh Thian / Ngaimin / Azmi", key="pms_operator")
        with c2:
            ms_rigger = st.text_input("Rigger / Signalman", "Rizal / Hanifah / Aziz / Jamari / Ahmad / Rahman / Malik", key="pms_rigger")
            ms_lifting_crew = st.text_area(
                "Crew / Responsibility Notes",
                value="MOM certified / qualified lifting supervisor, rigger, signalman and equipment operator shall be deployed according to the actual lifting equipment and project requirements.",
                height=80,
                key="pms_lifting_crew"
            )

    with st.expander("6. Site Photos, Site Plan & Automatic Drawings", expanded=True):
        site_photos = st.file_uploader(
            "Upload Site Photos",
            type=["png", "jpg", "jpeg"],
            accept_multiple_files=True,
            key="pms_site_photos"
        )
        photo_caption_text = st.text_area(
            "Photo Captions — one line per uploaded photo (optional)",
            height=100,
            key="pms_photo_captions",
            placeholder="Photo 1: Lorry crane set-up area\nPhoto 2: Roller shutter entrance\nPhoto 3: Final machine position"
        )
        site_plan_file = st.file_uploader(
            "Upload Site Plan / Customer Drawing (optional)",
            type=["pdf", "png", "jpg", "jpeg"],
            key="pms_site_plan"
        )
        st.markdown("**Automatic drawings to include**")
        st.caption("Only include the drawings useful for this job. The defaults below follow the site arrangement, but you can change them.")
        d1, d2, d3, d4 = st.columns(4)
        crane_selected_for_ms = any(x in equipment for x in ["Lorry loader / lorry crane", "Mobile crane"])
        elevated_for_ms = "Elevated" in ms_site_arrangement
        draw_top = d1.checkbox("Top View", value=crane_selected_for_ms, key="pms2_draw_top")
        draw_side = d2.checkbox("Side Elevation", value=elevated_for_ms, key="pms2_draw_side")
        draw_route = d3.checkbox("Movement Route", value=True, key="pms2_draw_route")
        draw_rigging = d4.checkbox("Rigging Sketch", value=crane_selected_for_ms, key="pms2_draw_rigging")

    gear_items = []
    worker_items = []
    github_ready = bool(github_settings().get("token") and github_settings().get("repo"))

    with st.expander("7. Certificate / Supporting Document Appendices", expanded=False):
        st.caption("Select certificates already stored in your GitHub folders, or upload other supporting documents.")
        if github_ready:
            try:
                gear_items = github_list_folder("Lifting Gears Certificate")
                worker_items = github_list_folder("Workers Certificate")
            except Exception as exc:
                st.warning(f"Could not load GitHub certificate lists: {exc}")
        else:
            st.info("GitHub Secrets are not configured yet. You can still generate the MOS and upload supporting files manually.")

        gear_by_name = {i.get("name", ""): i for i in gear_items}
        worker_by_name = {i.get("name", ""): i for i in worker_items}
        selected_gear_names = st.multiselect(
            "Lifting Gear / Crane Certificates",
            list(gear_by_name.keys()),
            key="pms_selected_gear"
        )
        selected_worker_names = st.multiselect(
            "Worker Training Certificates",
            list(worker_by_name.keys()),
            key="pms_selected_workers"
        )
        extra_attachments = st.file_uploader(
            "Other Supporting Documents",
            type=["pdf", "png", "jpg", "jpeg", "docx"],
            accept_multiple_files=True,
            key="pms_extra_attachments"
        )
        embed_full_pdfs = st.checkbox(
            "Embed all pages of selected PDFs into the Word MOS",
            value=True,
            key="pms_embed_full_pdf"
        )
        st.caption("For PDF pages to be embedded, add `PyMuPDF` to requirements.txt. If it is missing, the MOS still generates but only lists the PDF attachment.")

    generate_ms = st.button("🏗️ Generate Professional MOS Package", key="generate_professional_ms", type="primary")

    if generate_ms:
        required = {
            "Customer / Tenant Company": ms_customer,
            "Project Name": ms_project_name,
            "Site Location": ms_location,
            "Description of Work": ms_description,
            "Machine / Load Name": ms_machine_name,
        }
        missing_fields = [label for label, value in required.items() if not str(value).strip()]
        if missing_fields:
            st.error("Please complete: " + ", ".join(missing_fields))
        else:
            try:
                with st.spinner("Generating project-specific method, drawings and Word package..."):
                    data = {
                        "customer": ms_customer,
                        "project_name": ms_project_name,
                        "location": ms_location,
                        "description": ms_description,
                        "process": ms_process,
                        "operation_date": str(ms_date_input),
                        "operation_time": ms_operation_time,
                        "prepared_by": ms_prepared_by,
                        "assessed_by": ms_assessed_by,
                        "approved_by": ms_approved_by,
                        "last_review": str(ms_last_review),
                        "next_review": str(ms_next_review),
                        "machine_name": ms_machine_name,
                        "load_type": ms_load_type,
                        "load_weight_kg": float(ms_load_weight),
                        "rigging_weight_kg": float(ms_rigging_weight),
                        "total_lifted_kg": total_lifted,
                        "length_m": float(ms_length),
                        "width_m": float(ms_width),
                        "height_m": float(ms_height),
                        "cg": ms_cg,
                        "lifting_points": ms_lifting_points,
                        "sling_angle": float(ms_sling_angle),
                        "lift_classification": ms_lift_class,
                        "lift_reason": ms_lift_reason,
                        "equipment": equipment,
                        "crane_model": ms_crane_model,
                        "crane_lm": ms_crane_lm,
                        "crane_max_capacity": ms_crane_max_capacity,
                        "boom_length": float(ms_boom_length),
                        "operating_radius": float(ms_operating_radius),
                        "swl_at_radius_kg": float(ms_swl_radius),
                        "utilisation_pct": float(utilisation),
                        "forklift_details": ms_forklift_details,
                        "lifting_gear": ms_lifting_gear,
                        "site_arrangement": ms_site_arrangement,
                        "ground_condition": ms_ground,
                        "access_route": ms_access,
                        "obstacles": ms_obstacles,
                        "environment": ms_environment,
                        "building_height": float(ms_building_height),
                        "crane_distance": float(ms_crane_distance),
                        "landing_depth": float(ms_landing_depth),
                        "door_width": float(ms_door_width),
                        "door_height": float(ms_door_height),
                        "route_start": ms_route_start,
                        "route_final": ms_route_final,
                        "route_notes": ms_route_notes,
                        "lifting_supervisor": ms_lifting_supervisor,
                        "operator": ms_operator,
                        "rigger": ms_rigger,
                        "lifting_crew": ms_lifting_crew,
                        "draw_top": draw_top,
                        "draw_side": draw_side,
                        "draw_route": draw_route,
                        "draw_rigging": draw_rigging,
                    }

                    ai_content = mos_generate_ai_content(data)
                    captions = [line.strip() for line in photo_caption_text.splitlines() if line.strip()]

                    selected_cert_items = []
                    for name in selected_gear_names:
                        if name in gear_by_name:
                            selected_cert_items.append(("Lifting Equipment / Gear Certificate", gear_by_name[name]))
                    for name in selected_worker_names:
                        if name in worker_by_name:
                            selected_cert_items.append(("Worker Training Certificate", worker_by_name[name]))

                    buffer = mos_build_professional_doc(
                        data=data,
                        ai=ai_content,
                        site_photos=site_photos or [],
                        photo_captions=captions,
                        site_plan_file=site_plan_file,
                        cert_items=selected_cert_items,
                        extra_attachments=extra_attachments or [],
                        full_pdf=embed_full_pdfs,
                    )

                    safe_project = re.sub(r"[^A-Za-z0-9_-]+", "_", ms_project_name.strip())[:50] or "Project"
                    st.success("Professional Method Statement package generated.")
                    st.download_button(
                        "⬇️ Download Professional MOS (.docx)",
                        buffer,
                        file_name=f"EWMT_Professional_MOS_{safe_project}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )

                    with st.expander("Preview AI-generated work sequence"):
                        for i, step in enumerate(ai_content.get("work_sequence", []), start=1):
                            st.write(f"{i}. {step}")

            except Exception as e:
                st.error("Professional Method Statement generation failed")
                st.exception(e)


# ======================================================
# LIFTING PLAN
# ======================================================
if page == "🏗️ Lifting Plan":
    st.markdown("## 🏗️ Lifting Plan / Permit to Work")
    st.caption("Fill in all lifting details. Checkboxes selected here will be inserted into your Word template.")

    with st.expander("1. General", expanded=True):
        lp_company = st.text_input("Company", "Eric Wong Machinery Transportation Pte Ltd", key="lp_company")
        lp_project_name = st.text_input("Project", key="lp_project_name")
        lp_location = st.text_input("Location of Lifting Operation", key="lp_location")
        lp_date_input = st.date_input("Date", value=date.today(), key="lp_date_input")
        lp_operation_time = st.text_input("Operation Time", key="lp_operation_time")
        lp_validity = st.text_input("Validity Period of Lifting Operation", "1 Day", key="lp_validity")

    with st.expander("2. Details of Loads to be Hoist", expanded=True):
        lp_description = st.text_area("Description of Load", key="lp_description")
        lp_machine = st.text_input("Machine Name / Spec", key="lp_machine")
        lp_machine_dimension = st.text_input("Overall Dimension of Load", key="lp_machine_dimension")
        lp_machine_weight = st.text_input("Weight of Load", key="lp_machine_weight")

        weight_known = st.checkbox("Known Weight", value=True, key="weight_known")
        weight_estimated = st.checkbox("Estimated Weight", value=False, key="weight_estimated")

        cg_obvious = st.checkbox("Centre of Gravity - Obvious", value=True, key="cg_obvious")
        cg_estimated = st.checkbox("Centre of Gravity - Estimated", value=False, key="cg_estimated")
        cg_drawing = st.checkbox("Centre of Gravity - Determined by Drawing", value=False, key="cg_drawing")

    with st.expander("3. Details of Lifting Equipment", expanded=True):
        mobile_crane = st.checkbox("Mobile Crane", value=False, key="mobile_crane")
        lorry_loader = st.checkbox("Lorry Loader", value=True, key="lorry_loader")

        crane_name = st.text_input("LM / LE Registration No.", key="crane_name")
        crane_renew = st.text_input("Date of Last Certification", key="crane_renew")
        crane_expiry = st.text_input("Expiry Date of Certificate", key="crane_expiry")
        crane_swl = st.text_input("Max Safe Working Load", key="crane_swl")
        boom_length = st.text_input("Max Boom / Jib Length", key="boom_length")
        crane_radius = st.text_input("Intended Load Radius", key="crane_radius")
        crane_swl_radius = st.text_input("SWL at This Radius", key="crane_swl_radius")

        lifting_gear_manual = st.text_area(
            "Type of Lifting Gears",
            height=120,
            value="Wire rope slings / webbing slings, Shackles, Timber mats / steel plates, Tag lines",
            key="lifting_gear_manual"
        )

        lg_weight = st.text_input("Combined Weight of Lifting Gears", key="lg_weight")
        total_swl_lg = st.text_input("SWL of Lifting Gear", key="total_swl_lg")

        lg_cert_yes = st.checkbox("Certification of Lifting Gears - Yes", value=True, key="lg_cert_yes")
        lg_cert_no = st.checkbox("Certification of Lifting Gears - No", value=False, key="lg_cert_no")

        lg_expiry = st.text_input("Expiry Date of Lifting Gear Certificate", key="lg_expiry")

    with st.expander("4. Means of Communication", expanded=True):
        operator_can_see_yes = st.checkbox("Operator Can See Loading / Unloading Position - Yes", value=True, key="operator_can_see_yes")
        operator_can_see_no = st.checkbox("Operator Can See Loading / Unloading Position - No", value=False, key="operator_can_see_no")

        comm_standard = st.checkbox("Standard Hand Signals", value=True, key="comm_standard")
        comm_radio = st.checkbox("Radio", value=False, key="comm_radio")
        comm_others = st.checkbox("Others", value=False, key="comm_others")
        comm_others_text = st.text_input("Others Communication Details", key="comm_others_text")

    with st.expander("5. Personnel Involved in Lifting Operation", expanded=True):
        site_supervisor = st.text_input("Site Supervisor", "Ibrahim / Zahari / Zaharin / Wong Yen Siong", key="site_supervisor")
        lifting_supervisor = st.text_input("Lifting Supervisor", "Ibrahim / Zahari / Zaharin / Wong Yen Siong", key="lifting_supervisor")
        equipment_operator = st.text_input("Lifting Equipment Operator", "Lim Poh Soon / Norhalim / Lim Poh Thian / Ngaimin / Azmi", key="equipment_operator")
        rigger_1 = st.text_input("Rigger / Signalman 1", "Rizal / Hanifah / Aziz / Jamari / Ahmad", key="rigger_1")
        rigger_2 = st.text_input("Rigger / Signalman 2", "Rahman / Malik / Sarawanan / Sing Kwok Liang", key="rigger_2")

    with st.expander("6. Physical and Environmental Considerations", expanded=True):
        ground_safe_yes = st.checkbox("Ground Made Safe - Yes", value=True, key="ground_safe_yes")
        ground_safe_no = st.checkbox("Ground Made Safe - No", value=False, key="ground_safe_no")

        outriggers_yes = st.checkbox("Outriggers Evenly Extended - Yes", value=True, key="outriggers_yes")
        outriggers_no = st.checkbox("Outriggers Evenly Extended - No", value=False, key="outriggers_no")

        overhead_obstacle_yes = st.checkbox("Overhead Obstacles - Yes", value=False, key="overhead_obstacle_yes")
        overhead_obstacle_no = st.checkbox("Overhead Obstacles - No", value=True, key="overhead_obstacle_no")

        obstruction_yes = st.checkbox("Structure / Equipment / Materials Obstruction - Yes", value=False, key="obstruction_yes")
        obstruction_no = st.checkbox("Structure / Equipment / Materials Obstruction - No", value=True, key="obstruction_no")

        lighting_yes = st.checkbox("Lighting Adequate - Yes", value=True, key="lighting_yes")
        lighting_no = st.checkbox("Lighting Adequate - No", value=False, key="lighting_no")

        barricade_yes = st.checkbox("Zone Barricaded / Demarcated - Yes", value=True, key="barricade_yes")
        barricade_no = st.checkbox("Zone Barricaded / Demarcated - No", value=False, key="barricade_no")

        other_precautions = st.text_area("Other Precautions", key="other_precautions")

    with st.expander("7. Tasks", expanded=True):
        task_sequence = st.text_area(
            "Sequence of Lifting Operations",
            height=280,
            value="""1. Deploy lorry loader at designated unloading area
2. Set up crane with outriggers fully extended and resting on timber mats as base plate
3. Rigger to insert sling to crane hook
4. Secure sling to rigging point of load
5. Lorry loader to hoist down load from lorry chassis to the ground
6. Using forklift to unload and fork down machine from lorry chassis to the ground
7. Transport machine to door entrance
8. Using pallet truck to shift and position machine into factory premise
9. Position at the designated location
10. Once job complete, carry out proper housekeeping
11. All debris will be cleared and disposed
12. Job complete""",
            key="task_sequence"
        )

        person_in_charge = st.text_input(
            "Person in Charge for Each Step",
            "Zahari / Ibrahim / Wong Yen Siong",
            key="person_in_charge"
        )

    with st.expander("8. Approval of Lifting Plan", expanded=True):
        applied_by = st.text_input("Applied By", "Zailani", key="applied_by")
        applied_designation = st.text_input("Applied By Designation", "Supervisor", key="applied_designation")
        prepared_by = st.text_input("Prepared By", "Zahari", key="prepared_by_lp")
        prepared_designation = st.text_input("Prepared By Designation", "Lifting Supervisor", key="prepared_designation")
        assessed_by = st.text_input("Assessed By", "Kevin Wong", key="assessed_by")
        assessed_designation = st.text_input("Assessed By Designation", "Project Manager", key="assessed_designation")
        approved_by = st.text_input("Approved By", "Eric Wong", key="approved_by")
        approved_designation = st.text_input("Approved By Designation", "Managing Director", key="approved_designation")

    generate_lp = st.button("🏗️ Generate Lifting Plan", key="generate_lp")

    if generate_lp:
        try:
            with st.spinner("Generating Lifting Plan..."):

                response = client.responses.create(
                    model="gpt-5.4",
                    input=f"""
Improve this lifting task sequence into formal lifting plan wording.

Task sequence:
{task_sequence}

Return JSON only:
{{
 "lifting_method": "",
 "safety_controls": ""
}}
""",
                    tools=[{
                        "type": "file_search",
                        "vector_store_ids": [LP_VECTOR_STORE_ID]
                    }],
                    text={
                        "format": {
                            "type": "json_schema",
                            "name": "lp_schema",
                            "schema": {
                                "type": "object",
                                "additionalProperties": False,
                                "properties": {
                                    "lifting_method": {"type": "string"},
                                    "safety_controls": {"type": "string"}
                                },
                                "required": ["lifting_method", "safety_controls"]
                            }
                        }
                    }
                )

                data = json.loads(response.output_text)
                doc = Document(LP_TEMPLATE)

                replacements = {
                    "{{company}}": safe_text(lp_company),
                    "{{project_name}}": safe_text(lp_project_name),
                    "{{location}}": safe_text(lp_location),
                    "{{date}}": str(lp_date_input),
                    "{{operation_date}}": str(lp_date_input),
                    "{{operation_time}}": safe_text(lp_operation_time),
                    "{{validity_period}}": safe_text(lp_validity),

                    "{{description_of_work}}": safe_text(lp_description),
                    "{{machine_spec}}": safe_text(lp_machine),
                    "{{machine_name}}": safe_text(lp_machine),
                    "{{machine_dimension}}": safe_text(lp_machine_dimension),
                    "{{machine_weight}}": safe_text(lp_machine_weight),

                    "{{kw}}": tick(weight_known),
                    "{{ew}}": tick(weight_estimated),
                    "{{obv}}": tick(cg_obvious),
                    "{{Est}}": tick(cg_estimated),
                    "{{est}}": tick(cg_estimated),
                    "{{ddw}}": tick(cg_drawing),

                    "{{mob_cr}}": tick(mobile_crane),
                    "{{lor_cr}}": tick(lorry_loader),

                    "{{Crane_lm}}": safe_text(crane_name),
                    "{{crane_lm}}": safe_text(crane_name),
                    "{{crane_name}}": safe_text(crane_name),
                    "{{crane_renew}}": safe_text(crane_renew),
                    "{{crane_expiry}}": safe_text(crane_expiry),
                    "{{crane_swl}}": safe_text(crane_swl),
                    "{{boom_length}}": safe_text(boom_length),

                    "{{crane_radius}}": safe_text(crane_radius),
                    "{{ crane_radius }}": safe_text(crane_radius),
                    "{{crane_radius }}": safe_text(crane_radius),
                    "{{ crane_radius}}": safe_text(crane_radius),
                    "{{crane_swl_radius}}": safe_text(crane_swl_radius),

                    "{{lifting_gear}}": safe_text(lifting_gear_manual),
                    "{{lg_weight}}": safe_text(lg_weight),
                    "{{lifting_gear_wt}}": safe_text(lg_weight),
                    "{{total_swl_lg}}": safe_text(total_swl_lg),

                    "{{c_lg_y}}": tick(lg_cert_yes),
                    "{{c_lg_n}}": tick(lg_cert_no),
                    "{{lg_expiry}}": safe_text(lg_expiry),

                    "{{coms_y}}": tick(operator_can_see_yes),
                    "{{coms_n}}": tick(operator_can_see_no),
                    "{{coms}}": tick(operator_can_see_yes),

                    "{{shs}}": tick(comm_standard),
                    "{{rad}}": tick(comm_radio),
                    "{{comm_standard}}": tick(comm_standard),
                    "{{comm_radio}}": tick(comm_radio),
                    "{{comm_others}}": tick(comm_others),
                    "{{comm_others_text}}": safe_text(comm_others_text),

                    "{{site_supervisor}}": safe_text(site_supervisor),
                    "{{lifting_supervisor}}": safe_text(lifting_supervisor),
                    "{{equipment_operator}}": safe_text(equipment_operator),
                    "{{rigger_1}}": safe_text(rigger_1),
                    "{{rigger_2}}": safe_text(rigger_2),

                    "{{gc_y}}": tick(ground_safe_yes),
                    "{{gc_n}}": tick(ground_safe_no),
                    "{{go_y}}": tick(outriggers_yes),
                    "{{go_n}}": tick(outriggers_no),
                    "{{ob_y}}": tick(overhead_obstacle_yes),
                    "{{ob_n}}": tick(overhead_obstacle_no),
                    "{{st_y}}": tick(obstruction_yes),
                    "{{st_n}}": tick(obstruction_no),
                    "{{li_y}}": tick(lighting_yes),
                    "{{li_n}}": tick(lighting_no),
                    "{{de_y}}": tick(barricade_yes),
                    "{{de_n}}": tick(barricade_no),
                    "{{other_precautions}}": safe_text(other_precautions),

                    "{{task_sequence}}": safe_text(task_sequence),
                    "{{tasks}}": safe_text(task_sequence),
                    "{{lifting_method}}": safe_text(data.get("lifting_method", "")),
                    "{{safety_controls}}": safe_text(data.get("safety_controls", "")),
                    "{{person_in_charge}}": safe_text(person_in_charge),
                    "{{task_pic}}": safe_text(person_in_charge),

                    "{{applied_by}}": safe_text(applied_by),
                    "{{applied_designation}}": safe_text(applied_designation),
                    "{{prepared_by}}": safe_text(prepared_by),
                    "{{prepared_designation}}": safe_text(prepared_designation),
                    "{{assessed_by}}": safe_text(assessed_by),
                    "{{assessed_designation}}": safe_text(assessed_designation),
                    "{{approved_by}}": safe_text(approved_by),
                    "{{approved_designation}}": safe_text(approved_designation),

                    "{{known_weight_checked}}": tick(weight_known),
                    "{{estimated_weight_checked}}": tick(weight_estimated),
                    "{{center_gravity_obvious}}": tick(cg_obvious),
                    "{{center_gravity_estimated}}": tick(cg_estimated),
                    "{{center_gravity_drawing}}": tick(cg_drawing),
                    "{{mobile_crane_checked}}": tick(mobile_crane),
                    "{{lorry_loader_checked}}": tick(lorry_loader),
                    "{{lg_cert_yes}}": tick(lg_cert_yes),
                    "{{lg_cert_no}}": tick(lg_cert_no),
                    "{{operator_can_see_yes}}": tick(operator_can_see_yes),
                    "{{operator_can_see_no}}": tick(operator_can_see_no),
                    "{{ground_safe_yes}}": tick(ground_safe_yes),
                    "{{ground_safe_no}}": tick(ground_safe_no),
                    "{{outriggers_yes}}": tick(outriggers_yes),
                    "{{outriggers_no}}": tick(outriggers_no),
                    "{{obstacles_yes}}": tick(overhead_obstacle_yes),
                    "{{obstacles_no}}": tick(overhead_obstacle_no),
                    "{{obstruction_yes}}": tick(obstruction_yes),
                    "{{obstruction_no}}": tick(obstruction_no),
                    "{{lighting_yes}}": tick(lighting_yes),
                    "{{lighting_no}}": tick(lighting_no),
                    "{{barricade_yes}}": tick(barricade_yes),
                    "{{barricade_no}}": tick(barricade_no),
                }

                replace_all(doc, replacements)

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.download_button(
                    "Download Lifting Plan",
                    buffer,
                    "Lifting_Plan.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        except Exception as e:
            st.error("Lifting Plan generation failed")
            st.exception(e)


# ======================================================
# RISK ASSESSMENT PRO
# ======================================================
if page == "⚠️ Risk Assessment Pro":
    st.markdown("## ⚠️ Risk Assessment Pro")
    st.caption("Generate a professional 5x5 Risk Assessment based on actual work activities.")

    with st.expander("Project Details", expanded=True):
        ra_company = st.text_input("Company", "Eric Wong Machinery Transportation Pte Ltd", key="ra_company")
        ra_project_name = st.text_input("Project Name", key="ra_project_name")
        ra_location = st.text_input("Location", key="ra_location")
        ra_machine = st.text_input("Machine Spec", key="ra_machine")
        ra_description = st.text_area("Description of Work", key="ra_description")
        ra_date_input = st.date_input("Date", value=date.today(), key="ra_date_input")
        ra_due_date_input = st.date_input("Due Date", value=date.today(), key="ra_due_date_input")

    with st.expander("Risk Assessment Details", expanded=True):
        ra_process = st.text_input("RA Process", "Machinery Moving / Lifting Operation", key="ra_process")

        activities = st.text_area(
            "Work Activities (1 per line)",
            height=220,
            value="""Transport of lifting machinery into or out of site premises
Setting up of crane on site
Lifting operation
Signalling of load""",
            key="activities"
        )

    generate_ra_pro = st.button("⚠️ Generate Risk Assessment Pro", key="generate_ra_pro")

    if generate_ra_pro:
        try:
            with st.spinner("Generating Risk Assessment Pro..."):

                prompt = f"""
Create professional Singapore style 5x5 Risk Assessment.

Company: {ra_company}
Project: {ra_project_name}
Location: {ra_location}
Process: {ra_process}
Machine: {ra_machine}
Description: {ra_description}
Date: {ra_date_input}
Due Date: {ra_due_date_input}

Activities:
{activities}

Important:
- Every generated row must directly match the work activities provided.
- Do not invent unrelated activities.
- For each activity, create 1 to 3 relevant hazards.
- If one activity has more than one hazard, only the first row should contain ref and work_activity.
- Subsequent hazard rows for the same activity must use empty string for ref and work_activity.
- Use machinery moving / lifting / forklift / jacking / roller / crating style controls.
- Use wording style similar to Eric Wong Machinery Transportation Pte Ltd RA examples.
- Return JSON only.

Schema:
{{
 "rows":[
   {{
    "ref":"1",
    "work_activity":"",
    "hazard":"",
    "possible_injury":"",
    "existing_controls":"",
    "s":"4",
    "l":"2",
    "rpn":"8",
    "additional_controls":"",
    "rs":"4",
    "rl":"1",
    "rrpn":"4",
    "person":"Supervisor on site",
    "due_date":"{ra_due_date_input}",
    "remark":""
   }}
 ]
}}
"""

                response = client.responses.create(
                    model="gpt-5.4",
                    input=prompt,
                    tools=[{
                        "type": "file_search",
                        "vector_store_ids": [RA_VECTOR_STORE_ID]
                    }]
                )

                data = json.loads(response.output_text)

                doc = Document(RA_TEMPLATE)

                replace_all(doc, {
                    "{{company}}": ra_company,
                    "{{location}}": ra_location,
                    "{{process}}": ra_process,
                    "{{date}}": str(ra_date_input),
                    "{{due_date}}": str(ra_due_date_input)
                })

                fill_inventory_table(doc, activities, ra_location, ra_process)

                table = find_ra_table(doc)

                if table:
                    clear_rows_after_column_header(table)

                    for r in data["rows"]:
                        add_ra_row(table, [
                            r["ref"],
                            r["work_activity"],
                            r["hazard"],
                            r["possible_injury"],
                            r["existing_controls"],
                            r["s"],
                            r["l"],
                            r["rpn"],
                            r["additional_controls"],
                            r["rs"],
                            r["rl"],
                            r["rrpn"],
                            r["person"],
                            str(ra_due_date_input),
                            r["remark"]
                        ])

                    merge_same_work_activity_cells(table)

                format_risk_assessment(doc)

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.download_button(
                    "Download Risk Assessment Pro",
                    buffer,
                    "Risk_Assessment_Pro.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        except Exception as e:
            st.error("Risk Assessment Pro generation failed")
            st.exception(e)


# ======================================================
# LIFTING GEAR REGISTER
# ======================================================
if page == "🧰 Lifting Gear Register":
    certificate_browser(
        folder_name="Lifting Gears Certificate",
        title="🧰 Lifting Gear Register",
        info_text="Certificates loaded from GitHub folder: Lifting Gears Certificate",
        search_label="Search by SWL / keyword",
        search_placeholder="Example: 3 Ton, 10 Ton, shackle, round sling",
        download_label="Download Selected Certificate"
    )


# ======================================================
# WORKER TRAINING CERTIFICATE
# ======================================================
if page == "👷 Worker Training Certificate":
    certificate_browser(
        folder_name="Workers Certificate",
        title="👷 Worker Training Certificate",
        info_text="Certificates loaded from GitHub folder: Workers Certificate",
        search_label="Search by worker name / course / keyword",
        search_placeholder="Example: Ibrahim, forklift, boom lift, lifting supervisor, rigger",
        download_label="Download Selected Worker Certificate"
    )


# ======================================================
# EXPIRY ALERTS
# ======================================================
if page == "⏰ Expiry Alerts":
    st.markdown("## ⏰ Expiry Alerts")
    st.caption("Show expired and expiring lifting gear certificates from your GitHub folder.")

    import re

    CERT_FOLDER = os.path.join(BASE_DIR, "Lifting Gears Certificate")

    if not os.path.exists(CERT_FOLDER):
        st.error("Folder not found: Lifting Gears Certificate")
        st.code("Lifting Gears Certificate")
    else:
        files = [
            f for f in os.listdir(CERT_FOLDER)
            if f.lower().endswith((".pdf", ".png", ".jpg", ".jpeg"))
        ]

        if not files:
            st.warning("No certificate files found.")
        else:
            today = date.today()
            alert_days = st.number_input(
                "Show certificates expiring within how many days?",
                min_value=1,
                max_value=365,
                value=30
            )

            records = []

            for f in files:
                found_date = None

                patterns = [
                    r"(\d{4})[-_\\.](\d{1,2})[-_\\.](\d{1,2})",
                    r"(\d{1,2})[-_\\.](\d{1,2})[-_\\.](\d{4})",
                ]

                for pattern in patterns:
                    match = re.search(pattern, f)
                    if match:
                        try:
                            parts = match.groups()

                            if len(parts[0]) == 4:
                                found_date = date(int(parts[0]), int(parts[1]), int(parts[2]))
                            else:
                                found_date = date(int(parts[2]), int(parts[1]), int(parts[0]))

                            break
                        except Exception:
                            found_date = None

                if found_date:
                    days_left = (found_date - today).days

                    if days_left < 0:
                        status = "Expired"
                    elif days_left <= alert_days:
                        status = "Expiring Soon"
                    else:
                        status = "Valid"

                    records.append({
                        "Certificate File": f,
                        "Expiry Date": str(found_date),
                        "Days Left": days_left,
                        "Status": status
                    })
                else:
                    records.append({
                        "Certificate File": f,
                        "Expiry Date": "No date found in filename",
                        "Days Left": "",
                        "Status": "Unknown"
                    })

            expired = [r for r in records if r["Status"] == "Expired"]
            expiring = [r for r in records if r["Status"] == "Expiring Soon"]
            valid = [r for r in records if r["Status"] == "Valid"]

            c1, c2, c3 = st.columns(3)
            c1.metric("Expired", len(expired))
            c2.metric("Expiring Soon", len(expiring))
            c3.metric("Valid", len(valid))

            st.markdown("### Certificate Expiry List")
            st.dataframe(records, use_container_width=True)

            st.info("For this to work, put expiry date inside the certificate filename, example: 3 Ton Shackle Expiry 2026-06-30.pdf")


# ======================================================
# SETTINGS / ADMIN DOCUMENT MANAGER
# ======================================================
if page == "⚙️ Settings":
    st.markdown("## ⚙️ Settings / Administrator")
    st.caption("Secure document control and template reference for the EWMT internal system.")

    if not admin_is_logged_in():
        admin_login_form()
    else:
        render_admin_document_manager()

    st.markdown("---")
    with st.expander("📄 Method Statement Placeholder Guide"):
        st.code("""
Use these placeholders in your Method Statement Word template:

{{date}}
{{description_of_work}}
{{machine_spec}}
{{operation_date}}
{{operation_time}}
{{location}}
{{equipment}}
{{obstacles}}
{{environment}}
{{lifting_crew}}
{{safety_aspect}}
{{job_scope}}
{{prepared_by}}
""")

    with st.expander("🏗️ Lifting Plan Placeholder Guide"):
        st.code("""
Use these placeholders in your Lifting Plan Word template:

General:
{{project_name}}
{{location}}
{{operation_date}}
{{operation_time}}
{{validity_period}}

Load:
{{machine_name}}
{{machine_dimension}}
{{machine_weight}}
{{kw}} Known weight
{{ew}} Estimated weight
{{obv}} Obvious
{{Est}} Estimated
{{ddw}} Determined by drawing

Lifting Equipment:
{{mob_cr}} Mobile crane
{{lor_cr}} Lorry loader
{{Crane_lm}}
{{crane_lm}}
{{crane_renew}}
{{crane_expiry}}
{{crane_swl}}
{{crane_radius}}
{{crane_swl_radius}}
{{lifting_gear}}
{{lifting_gear_wt}}
{{total_swl_lg}}
{{c_lg_y}} Yes
{{c_lg_n}} No
{{lg_expiry}}

Communication:
{{coms_y}} Yes
{{coms_n}} No
{{shs}} Standard hand signals
{{rad}} Radio

Physical / Environmental:
{{gc_y}} Yes    {{gc_n}} No
{{go_y}} Yes    {{go_n}} No
{{ob_y}} Yes    {{ob_n}} No
{{st_y}} Yes    {{st_n}} No
{{li_y}} Yes    {{li_n}} No
{{de_y}} Yes    {{de_n}} No

Approval:
{{applied_by}}
{{applied_designation}}
{{prepared_by}}
{{prepared_designation}}
{{assessed_by}}
{{assessed_designation}}
{{approved_by}}
{{approved_designation}}
""")
